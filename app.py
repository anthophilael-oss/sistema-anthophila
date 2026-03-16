import streamlit as st
import pandas as pd
import json
import os
import urllib.parse
from datetime import datetime, timedelta
from streamlit_calendar import calendar
from PIL import Image
from docx import Document
from docx.shared import Pt

# Intentar importar librerías opcionales para facturación
try:
    from reportlab.lib.pagesizes import A4, A5
    from reportlab.pdfgen import canvas
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    import qrcode
    import barcode
    from barcode.writer import ImageWriter
    HAS_BILLING_LIBS = True
except ImportError:
    HAS_BILLING_LIBS = False

# --- CONFIGURACIÓN E IDENTIDAD ---
st.set_page_config(page_title="Anthophila v9.9.5", layout="wide", page_icon="🐝")

# Lista de servicios predefinidos
SERVICIOS_DISPONIBLES = [
    "Consulta Psicológica / Terapia",
    "Evaluación Neuropsicológica",
    "Informe Psicológico / Certificado",
    "Programa de Lecto-Escritura",
    "Programa de Desarrollo Cognitivo",
    "Terapia de Lenguaje",
    "Terapia de Aprendizaje",
    "Asesoría a Padres / Orientación",
    "Sesión de Evaluación Inicial",
    "Otro (Especificar)"
]

# Catálogo SUNAT No. 09: Motivos de emisión de Nota de Crédito
MOTIVOS_NOTA_CREDITO = [
    "01 - Anulación de la operación",
    "02 - Anulación por error en el RUC",
    "03 - Corrección por error en la descripción",
    "04 - Descuento global",
    "05 - Descuento por ítem",
    "06 - Devolución total",
    "07 - Devolución por ítem",
    "08 - Bonificación",
    "09 - Disminución en el valor",
    "10 - Otros conceptos"
]

# --- ESTILOS PERSONALIZADOS (CSS) ---
st.markdown("""
    <style>
    /* Fondo general */
    .stApp {
        background-color: #f8fbf8;
    }
    
    /* Estilo para los botones principales */
    .stButton>button {
        background-color: #27AE60;
        color: white;
        border-radius: 8px;
        border: none;
        padding: 0.6rem 1.2rem;
        font-weight: bold;
        transition: all 0.3s ease;
        width: 100%;
    }
    .stButton>button:hover {
        background-color: #1E8449;
        box-shadow: 0 4px 8px rgba(0,0,0,0.1);
        transform: translateY(-2px);
    }
    
    /* Estilo para los inputs */
    .stTextInput>div>div>input {
        border-radius: 8px;
    }
    
    /* Contenedores con sombra suave */
    div[data-testid="stVerticalBlock"] > div[style*="border"] {
        background-color: white;
        padding: 2rem;
        border-radius: 12px;
        box-shadow: 0 4px 12px rgba(0,0,0,0.05);
        border: 1px solid #e0e0e0 !important;
    }
    
    /* Títulos con color temático */
    h1, h2, h3 {
        color: #1B5E20;
        font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
    }
    
    /* Sidebar estilizado */
    [data-testid="stSidebar"] {
        background-color: #ffffff;
        border-right: 1px solid #eee;
    }
    
    /* Tabs estilizadas */
    .stTabs [data-baseweb="tab-list"] {
        gap: 10px;
    }
    .stTabs [data-baseweb="tab"] {
        background-color: #f1f1f1;
        border-radius: 8px 8px 0 0;
        padding: 10px 20px;
    }
    .stTabs [aria-selected="true"] {
        background-color: #27AE60 !important;
        color: white !important;
    }
    </style>
""", unsafe_allow_html=True)

# Rutas de Archivos (Optimizadas para Streamlit Cloud / Rutas Relativas)
BASE_DIR = "PACIENTES_ANTHOPHILA"
DATA_FILE = "mapeo_dni.json"
CITAS_FILE = "citas_anthophila.json"
EXCEL_CITAS = "db_reservas_anthophila.csv"
ASIGNACIONES_FILE = "asignaciones_pruebas.json"
INGRESOS_FILE = "db_ingresos_anthophila.csv"
LIBRO_DIARIO_FILE = "libro_diario_5_2.csv"
PLAN_CONTABLE_FILE = "plan_contable_5_4.csv"
REGISTRO_COMPRAS_8_1_FILE = "registro_compras_8_1.csv"
REGISTRO_COMPRAS_8_2_FILE = "registro_compras_8_2.csv"
REGISTRO_VENTAS_14_1_FILE = "registro_ventas_14_1.csv"
DB_CLIENTES_FILE = "db_clientes_anthophila.csv"
IMG_FOLDER = "imagenes"
LOGO_FILE = "logo4_antho.jpg"
COMPROBANTES_FOLDER = "comprobantes"

# Crear carpetas necesarias si no existen
for folder in [BASE_DIR, IMG_FOLDER, COMPROBANTES_FOLDER]:
    if not os.path.exists(folder):
        os.makedirs(folder)

# Datos del Especialista
CELULAR_ESPECIALISTA = "51922768805" 
PSICOLOGO_FULL = "Jose Miguel Granda Cortez. Psicólogo Colegiado y Habilitado (C.Ps.P. 18983). Magíster en Problemas de Aprendizaje con estudios de Doctorado en Psicología Educativa"
USUARIO_ADMIN = "admin"
PASSWORD_ADMIN = "anthophila2026"
PASSWORD_FAMILIA = "familia2026"

# Información de Medios de Pago
INFO_PAGOS = {
    "Yape": {
        "detalle": "Celular: 922768805 (A nombre de Jose Miguel Granda)",
        "logo": "logo yape.jfif"
    },
    "Caja Tacna": {
        "detalle": "A nombre de: Anthophila E-Learning\nNº de Cuenta: 009211000130688\nCCI: 81300921100013068880",
        "logo": "logo caja.png"
    }
}

# Credenciales SUNAT Reales
SUNAT_RUC = "20607755907"
SUNAT_USER = "43707088"
SUNAT_PASS = "Jj189114125"
CERT_PATH = "certificado.p12"

# --- FUNCIONES DE PERSISTENCIA ---
def cargar_json(archivo):
    if os.path.exists(archivo):
        with open(archivo, "r", encoding='utf-8') as f:
            try: return json.load(f)
            except: return [] if "citas" in archivo else {}
    return [] if "citas" in archivo else {}

def guardar_json(archivo, datos):
    with open(archivo, "w", encoding='utf-8') as f:
        json.dump(datos, f, ensure_ascii=False, indent=4)

def registrar_en_excel(datos):
    df_nueva = pd.DataFrame([datos])
    if not os.path.isfile(EXCEL_CITAS):
        df_nueva.to_csv(EXCEL_CITAS, index=False, sep=';', encoding='utf-8-sig')
    else:
        df_nueva.to_csv(EXCEL_CITAS, mode='a', index=False, header=False, sep=';', encoding='utf-8-sig')

def generar_asiento_contable(fec_str, monto, cliente, observacion, es_salida=False):
    """Genera un asiento en el Libro Diario 5.2 (Simplificado)"""
    periodo = datetime.strptime(fec_str, "%Y-%m-%d").strftime("%Y%m00")
    cuo = f"CUO-{datetime.now().strftime('%Y%m%d%H%M%S')}"
    
    asientos = []
    
    # Lógica de cuentas simplificada (Ejemplo)
    # Ingreso: 12 (Cuentas por cobrar) vs 70 (Ventas)
    # Salida: 63 (Gastos) vs 42 (Cuentas por pagar)
    
    if not es_salida:
        # DEBE (1212)
        asientos.append({
            "Periodo": periodo, "CUO": cuo, "Asiento": "M001", "Cuenta": "1212", 
            "Moneda": "PEN", "Tipo_Doc": "00", "Num_Doc": "0", "Fec_Emi": fec_str,
            "Glosa": f"Ingreso: {cliente} - {observacion}", "Debe": monto, "Haber": 0.00, "Estado": "1"
        })
        # HABER (7011)
        asientos.append({
            "Periodo": periodo, "CUO": cuo, "Asiento": "M001", "Cuenta": "7011", 
            "Moneda": "PEN", "Tipo_Doc": "00", "Num_Doc": "0", "Fec_Emi": fec_str,
            "Glosa": f"Venta: {cliente}", "Debe": 0.00, "Haber": monto, "Estado": "1"
        })
    else:
        # DEBE (6391)
        asientos.append({
            "Periodo": periodo, "CUO": cuo, "Asiento": "M001", "Cuenta": "6391", 
            "Moneda": "PEN", "Tipo_Doc": "00", "Num_Doc": "0", "Fec_Emi": fec_str,
            "Glosa": f"Gasto: {cliente} - {observacion}", "Debe": monto, "Haber": 0.00, "Estado": "1"
        })
        # HABER (4212)
        asientos.append({
            "Periodo": periodo, "CUO": cuo, "Asiento": "M001", "Cuenta": "4212", 
            "Moneda": "PEN", "Tipo_Doc": "00", "Num_Doc": "0", "Fec_Emi": fec_str,
            "Glosa": f"Pago Gasto: {cliente}", "Debe": 0.00, "Haber": monto, "Estado": "1"
        })

    df_asientos = pd.DataFrame(asientos)
    if not os.path.isfile(LIBRO_DIARIO_FILE):
        df_asientos.to_csv(LIBRO_DIARIO_FILE, index=False, sep=';', encoding='utf-8-sig')
    else:
        df_asientos.to_csv(LIBRO_DIARIO_FILE, mode='a', index=False, header=False, sep=';', encoding='utf-8-sig')

def generar_registro_compra(fec_str, monto, proveedor, ruc, obs):
    """Genera una entrada en el Registro de Compras 8.1 (Simplificado)"""
    periodo = datetime.strptime(fec_str, "%Y-%m-%d").strftime("%Y%m00")
    cuo = f"CUO-RC-{datetime.now().strftime('%Y%m%d%H%M%S')}"
    
    # Cálculos base (IGV 18%)
    base_imponible = round(monto / 1.18, 2)
    igv = round(monto - base_imponible, 2)
    
    compra = {
        "Periodo": periodo, "CUO": cuo, "Asiento": "M001", "Fecha_Emi": fec_str,
        "Fecha_Venc": "", "Tipo_CP": "01", "Serie": "F001", "Anio_DUA": "",
        "Num_CP": "00001", "Num_Final": "", "Tipo_Doc_Prov": "6", "RUC": ruc if ruc else "00000000000",
        "Nombre_Prov": proveedor, "Base_Gravada": base_imponible, "IGV": igv,
        "Base_Mix": 0.0, "IGV_Mix": 0.0, "Base_NoGrav": 0.0, "IGV_NoGrav": 0.0,
        "NoGrav": 0.0, "ISC": 0.0, "ICBPER": 0.0, "Otros": 0.0, "Total": monto,
        "Moneda": "PEN", "Tipo_Cambio": 1.0, "Fec_Mod": "", "Tipo_Mod": "",
        "Serie_Mod": "", "Cod_Aduana": "", "Num_Mod": "", "Fec_Detrac": "",
        "Num_Detrac": "", "Marca_Ret": "", "Clasif": "", "Contrato": "",
        "Err1": "", "Err2": "", "Err3": "", "Err4": "", "Ind_Pago": "1", "Estado": "1"
    }
    
    df_compra = pd.DataFrame([compra])
    if not os.path.isfile(REGISTRO_COMPRAS_8_1_FILE):
        df_compra.to_csv(REGISTRO_COMPRAS_8_1_FILE, index=False, sep=';', encoding='utf-8-sig')
    else:
        df_compra.to_csv(REGISTRO_COMPRAS_8_1_FILE, mode='a', index=False, header=False, sep=';', encoding='utf-8-sig')

def generar_registro_venta(fec_str, monto, cliente, dni, obs):
    """Genera una entrada en el Registro de Ventas 14.1 (Simplificado)"""
    periodo = datetime.strptime(fec_str, "%Y-%m-%d").strftime("%Y%m00")
    cuo = f"CUO-RV-{datetime.now().strftime('%Y%m%d%H%M%S')}"
    
    # Cálculos base (IGV 18%)
    base_imponible = round(monto / 1.18, 2)
    igv = round(monto - base_imponible, 2)
    
    venta = {
        "Periodo": periodo, "CUO": cuo, "Asiento": "M001", "Fecha_Emi": fec_str,
        "Fecha_Venc": "", "Tipo_CP": "03", "Serie": "B001", "Num_CP": "00001",
        "Num_Final": "", "Tipo_Doc_Cli": "1" if len(str(dni)) == 8 else "6", 
        "DNI_RUC": dni if dni else "00000000", "Nombre_Cli": cliente,
        "Valor_Exp": 0.0, "Base_Gravada": base_imponible, "Desc_Base": 0.0,
        "IGV": igv, "Desc_IGV": 0.0, "Exonerado": 0.0, "Inafecto": 0.0,
        "ISC": 0.0, "Base_Arroz": 0.0, "IGV_Arroz": 0.0, "ICBPER": 0.0,
        "Otros": 0.0, "Total": monto, "Moneda": "PEN", "Tipo_Cambio": 1.0,
        "Fec_Mod": "", "Tipo_Mod": "", "Serie_Mod": "", "Num_Mod": "",
        "Contrato": "", "Err1": "", "Ind_Pago": "1", "Estado": "1"
    }
    
    df_venta = pd.DataFrame([venta])
    if not os.path.isfile(REGISTRO_VENTAS_14_1_FILE):
        df_venta.to_csv(REGISTRO_VENTAS_14_1_FILE, index=False, sep=';', encoding='utf-8-sig')
    else:
        df_venta.to_csv(REGISTRO_VENTAS_14_1_FILE, mode='a', index=False, header=False, sep=';', encoding='utf-8-sig')

def exportar_a_docx(datos, filename):
    doc = Document()
    doc.add_heading('HISTORIA CLÍNICA - ANTHOPHILA', 0)
    
    for key, value in datos.items():
        p = doc.add_paragraph()
        p.add_run(f"{key}: ").bold = True
        p.add_run(str(value))
    
    doc.save(filename)
    return filename

def generar_pdf_comprobante(datos_comp, filename):
    """Genera el PDF real de la Boleta/Factura en formato A5 (Optimizado para logo, texto y recuadro)"""
    if not HAS_BILLING_LIBS:
        return None
    
    # Formato A5: 14.8cm x 21.0cm
    c = canvas.Canvas(filename, pagesize=A5)
    width, height = A5
    
    verde_oscuro = colors.HexColor("#1B5E20")
    verde_claro = colors.HexColor("#27AE60")
    gris_fondo = colors.HexColor("#F9F9F9")
    gris_borde = colors.HexColor("#E0E0E0")
    
    # 1. Cabecera: Logo a la IZQUIERDA (Ajustado para ganar espacio horizontal)
    logo_path = os.path.abspath(os.path.join(IMG_FOLDER, LOGO_FILE))
    if os.path.exists(logo_path):
        try:
            # Logo más compacto y a la izquierda
            c.drawImage(logo_path, 0.4*cm, height - 3.2*cm, width=2.2*cm, height=2.2*cm, preserveAspectRatio=True)
        except:
            c.setFont("Helvetica-Bold", 10)
            c.drawString(0.5*cm, height - 2.0*cm, "ANTHOPHILA")
    
    # Emisor (JUSTO AL LADO DEL LOGO, X=2.8cm para ganar espacio)
    c.setFillColor(verde_oscuro)
    c.setFont("Helvetica-Bold", 8.5)
    c.drawString(2.8*cm, height - 1.2*cm, "ANTHOPHILA E-LEARNING EDUCATION E.I.R.L.")
    
    c.setFillColor(verde_claro)
    c.setFont("Helvetica-Oblique", 6.5)
    c.drawString(2.8*cm, height - 1.6*cm, "Atención especializada en desarrollo cognitivo conducta y aprendizaje")
    
    c.setFillColor(colors.black)
    c.setFont("Helvetica", 6.5)
    c.drawString(2.8*cm, height - 2.1*cm, f"RUC: {SUNAT_RUC}")
    c.drawString(2.8*cm, height - 2.5*cm, "Asoc. Para Grande Mza. B Lote 15 - Tacna")
    c.drawString(2.8*cm, height - 2.9*cm, "Contacto: +51 906 598 622")

    # 2. Recuadro SUNAT (Derecha, ajustado para no chocar con el texto largo)
    c.setStrokeColor(verde_claro)
    c.setLineWidth(1.2)
    # X=width - 5.2cm (9.6cm), W=4.6cm -> Deja más espacio para el texto central
    c.roundRect(width - 5.2*cm, height - 3.8*cm, 4.6*cm, 2.4*cm, 6, stroke=1, fill=0)
    
    c.setFont("Helvetica-Bold", 10)
    c.drawCentredString(width - 2.9*cm, height - 2.0*cm, datos_comp['tipo'].upper())
    c.setFont("Helvetica-Bold", 9)
    c.drawCentredString(width - 2.9*cm, height - 2.7*cm, f"RUC: {SUNAT_RUC}")
    c.setFillColor(colors.red)
    c.setFont("Helvetica-Bold", 11)
    c.drawCentredString(width - 2.9*cm, height - 3.4*cm, f"{datos_comp['serie']} - {datos_comp['correlativo']}")
    c.setFillColor(colors.black)

    # Si es Nota de Crédito, mostrar documento de referencia
    if "NOTA DE CRÉDITO" in datos_comp['tipo'].upper() and 'doc_referencia' in datos_comp:
        c.setFont("Helvetica-Bold", 7)
        c.drawCentredString(width - 2.9*cm, height - 3.8*cm, "DOC. REF.:")
        c.setFont("Helvetica", 7)
        c.drawCentredString(width - 2.9*cm, height - 4.2*cm, datos_comp['doc_referencia'])
        c.setFont("Helvetica-Bold", 6)
        c.drawCentredString(width - 2.9*cm, height - 4.5*cm, "MOTIVO:")
        c.setFont("Helvetica", 6)
        c.drawCentredString(width - 2.9*cm, height - 4.8*cm, datos_comp.get('motivo_nc', '')[:35])

    # 3. Datos Cliente (Debajo de la cabecera)
    y_cliente = height - 5.8*cm
    if "NOTA DE CRÉDITO" in datos_comp['tipo'].upper():
        y_cliente = height - 6.5*cm # Bajamos un poco si hay datos de referencia
    
    c.setStrokeColor(gris_borde)
    c.setFillColor(gris_fondo)
    c.roundRect(0.8*cm, y_cliente - 2.0*cm, width - 1.6*cm, 2.0*cm, 4, stroke=1, fill=1)
    
    c.setFillColor(colors.black)
    c.setFont("Helvetica-Bold", 7)
    c.drawString(1.1*cm, y_cliente - 0.4*cm, "ADQUIRIENTE")
    c.setFont("Helvetica", 8)
    nombre_display = datos_comp['cliente'].upper()[:55]
    c.drawString(1.1*cm, y_cliente - 1.0*cm, f"Señor(es): {nombre_display}")
    c.drawString(1.1*cm, y_cliente - 1.6*cm, f"DNI/RUC: {datos_comp['doc_cliente']}")
    
    c.setFont("Helvetica-Bold", 7)
    c.drawString(width - 5.5*cm, y_cliente - 0.4*cm, "FECHA Y MONEDA")
    c.setFont("Helvetica", 8)
    c.drawString(width - 5.5*cm, y_cliente - 1.0*cm, f"Fecha Emisión: {datos_comp['fecha']}")
    c.drawString(width - 5.5*cm, y_cliente - 1.6*cm, "Moneda: SOLES (PEN)")

    # 4. Tabla
    y_tabla = y_cliente - 3.5*cm
    c.setFillColor(verde_oscuro)
    c.rect(0.8*cm, y_tabla, width - 1.6*cm, 0.5*cm, stroke=0, fill=1)
    
    c.setFillColor(colors.white)
    c.setFont("Helvetica-Bold", 7)
    c.drawString(1.1*cm, y_tabla + 0.15*cm, "DESCRIPCIÓN DEL SERVICIO")
    c.drawRightString(width - 3.5*cm, y_tabla + 0.15*cm, "CANT.")
    c.drawRightString(width - 2.2*cm, y_tabla + 0.15*cm, "P. UNIT")
    c.drawRightString(width - 1.1*cm, y_tabla + 0.15*cm, "TOTAL")
    
    c.setFillColor(colors.black)
    c.setFont("Helvetica", 8)
    y_item = y_tabla - 0.6*cm
    c.drawString(1.1*cm, y_item, datos_comp['descripcion'][:48])
    c.drawRightString(width - 3.5*cm, y_item, str(datos_comp['cantidad']))
    c.drawRightString(width - 2.2*cm, y_item, f"{datos_comp['total']/datos_comp['cantidad']:,.2f}")
    c.drawRightString(width - 1.1*cm, y_item, f"{datos_comp['total']:,.2f}")
    
    c.setStrokeColor(verde_claro)
    c.line(0.8*cm, y_item - 0.2*cm, width - 0.8*cm, y_item - 0.2*cm)

    # 5. Totales
    y_totales = y_item - 2.0*cm
    total = float(datos_comp['total'])
    base = round(total / 1.18, 2)
    igv = round(total - base, 2)
    
    c.setFont("Helvetica", 8)
    c.drawString(width - 5.0*cm, y_totales, "OP. GRAVADA:")
    c.drawRightString(width - 1.1*cm, y_totales, f"S/. {base:,.2f}")
    c.drawString(width - 5.0*cm, y_totales - 0.4*cm, "I.G.V. (18%):")
    c.drawRightString(width - 1.1*cm, y_totales - 0.4*cm, f"S/. {igv:,.2f}")
    
    c.setFillColor(verde_claro)
    c.rect(width - 5.2*cm, y_totales - 1.0*cm, 4.4*cm, 0.5*cm, stroke=0, fill=1)
    c.setFillColor(colors.white)
    c.setFont("Helvetica-Bold", 9)
    c.drawString(width - 5.0*cm, y_totales - 0.8*cm, "TOTAL A PAGAR:")
    c.drawRightString(width - 1.1*cm, y_totales - 0.8*cm, f"S/. {total:,.2f}")

    # 6. Pie
    y_pie = 0.8*cm
    tipo_doc_sunat = "03" if "Boleta" in datos_comp['tipo'] else "01"
    tipo_doc_cli = "1" if len(str(datos_comp['doc_cliente'])) == 8 else "6"
    qr_content = f"{SUNAT_RUC}|{tipo_doc_sunat}|{datos_comp['serie']}|{datos_comp['correlativo']}|{igv}|{total}|{datos_comp['fecha']}|{tipo_doc_cli}|{datos_comp['doc_cliente']}"
    
    try:
        qr = qrcode.make(qr_content)
        qr_img_path = os.path.abspath("temp_qr.png")
        qr.save(qr_img_path)
        c.drawImage(qr_img_path, 0.8*cm, y_pie, width=2.2*cm, height=2.2*cm)
        if os.path.exists(qr_img_path): os.remove(qr_img_path)
    except:
        pass
    
    c.setFillColor(colors.black)
    c.setFont("Helvetica-Bold", 6)
    c.drawString(3.2*cm, y_pie + 1.8*cm, "Representación impresa de la Factura/Boleta Electrónica")
    c.setFont("Helvetica", 6)
    c.drawString(3.2*cm, y_pie + 1.5*cm, "Consulte su comprobante en: https://play.google.com/store/apps/details?id=com.ppd.se&hl=es_US")
    c.setFont("Helvetica-Oblique", 6)
    c.drawString(3.2*cm, y_pie + 1.0*cm, "Anthophila: Sembrando hoy el aprendizaje de mañana. Gracias por su confianza.")
    
    c.save()
    return filename

def generar_pdf_ticket_termico(datos_comp, filename):
    """Genera un PDF en formato Ticket Térmico (80mm) con Logo centrado"""
    if not HAS_BILLING_LIBS: return None
    
    from reportlab.lib.pagesizes import mm
    ticket_width = 80 * mm
    ticket_height = 200 * mm
    
    c = canvas.Canvas(filename, pagesize=(ticket_width, ticket_height))
    w = ticket_width
    h = ticket_height
    
    # 1. Logo Centrado
    logo_path = os.path.abspath(os.path.join(IMG_FOLDER, LOGO_FILE))
    if os.path.exists(logo_path):
        try:
            # Logo centrado y bien posicionado
            c.drawImage(logo_path, w/2 - 1.5*cm, h - 3.2*cm, width=3.0*cm, height=3.0*cm, preserveAspectRatio=True)
        except:
            pass
    
    # 2. Datos Empresa
    c.setFont("Helvetica-Bold", 9)
    c.drawCentredString(w/2, h - 3.8*cm, "ANTHOPHILA E-LEARNING EDUCATION E.I.R.L.")
    c.setFont("Helvetica", 7)
    c.drawCentredString(w/2, h - 4.2*cm, f"RUC: {SUNAT_RUC}")
    c.drawCentredString(w/2, h - 4.5*cm, "Asoc. Para Grande Mza. B Lote 15")
    c.drawCentredString(w/2, h - 4.8*cm, "Cel: +51 906 598 622")
    
    c.setLineWidth(0.5)
    c.line(0.5*cm, h - 5.1*cm, w - 0.5*cm, h - 5.1*cm)
    
    # 3. Datos Comprobante
    c.setFont("Helvetica-Bold", 10)
    c.drawCentredString(w/2, h - 5.6*cm, datos_comp['tipo'].upper())
    c.drawCentredString(w/2, h - 6.1*cm, f"{datos_comp['serie']} - {datos_comp['correlativo']}")
    
    # Si es Nota de Crédito en Ticket
    if "NOTA DE CRÉDITO" in datos_comp['tipo'].upper():
        c.setFont("Helvetica", 7)
        c.drawCentredString(w/2, h - 6.5*cm, f"Doc. Ref: {datos_comp.get('doc_referencia', '')}")
        c.drawCentredString(w/2, h - 6.8*cm, f"Motivo: {datos_comp.get('motivo_nc', '')[:30]}")
        y_ini_datos = h - 7.5*cm
    else:
        y_ini_datos = h - 6.8*cm

    c.setFont("Helvetica", 8)
    c.drawString(0.5*cm, y_ini_datos, f"Fecha: {datos_comp['fecha']}")
    c.drawString(0.5*cm, y_ini_datos - 0.4*cm, f"Cliente: {datos_comp['cliente'][:28].upper()}")
    c.drawString(0.5*cm, y_ini_datos - 0.8*cm, f"DNI/RUC: {datos_comp['doc_cliente']}")
    
    c.line(0.5*cm, y_ini_datos - 1.1*cm, w - 0.5*cm, y_ini_datos - 1.1*cm)
    
    # 4. Detalle
    c.setFont("Helvetica-Bold", 8)
    c.drawString(0.5*cm, y_ini_datos - 1.6*cm, "CANT  DESCRIPCIÓN")
    c.drawRightString(w - 0.5*cm, y_ini_datos - 1.6*cm, "TOTAL")
    
    c.setFont("Helvetica", 8)
    y = y_ini_datos - 2.1*cm
    c.drawString(0.5*cm, y, f"{datos_comp['cantidad']}   {datos_comp['descripcion'][:22]}")
    c.drawRightString(w - 0.5*cm, y, f"{datos_comp['total']:,.2f}")
    
    c.line(0.5*cm, y - 0.3*cm, w - 0.5*cm, y - 0.3*cm)
    
    # 5. Totales
    y_tot = y - 1.2*cm
    total = float(datos_comp['total'])
    base = round(total / 1.18, 2)
    igv = round(total - base, 2)
    
    c.setFont("Helvetica", 8)
    c.drawString(w - 4.2*cm, y_tot, "OP. GRAVADA:")
    c.drawRightString(w - 0.5*cm, y_tot, f"S/. {base:,.2f}")
    c.drawString(w - 4.2*cm, y_tot - 0.4*cm, "I.G.V. (18%):")
    c.drawRightString(w - 0.5*cm, y_tot - 0.4*cm, f"S/. {igv:,.2f}")
    
    c.setFont("Helvetica-Bold", 10)
    c.drawString(w - 4.2*cm, y_tot - 1.0*cm, "TOTAL:")
    c.drawRightString(w - 0.5*cm, y_tot - 1.0*cm, f"S/. {total:,.2f}")
    
    # 6. QR SUNAT
    tipo_doc_sunat = "03" if "Boleta" in datos_comp['tipo'] else "01"
    tipo_doc_cli = "1" if len(str(datos_comp['doc_cliente'])) == 8 else "6"
    qr_content = f"{SUNAT_RUC}|{tipo_doc_sunat}|{datos_comp['serie']}|{datos_comp['correlativo']}|{igv}|{total}|{datos_comp['fecha']}|{tipo_doc_cli}|{datos_comp['doc_cliente']}"
    
    try:
        qr = qrcode.make(qr_content)
        qr_path = os.path.abspath("temp_qr_thermal.png")
        qr.save(qr_path)
        c.drawImage(qr_path, w/2 - 1.1*cm, y_tot - 3.8*cm, width=2.2*cm, height=2.2*cm)
        if os.path.exists(qr_path): os.remove(qr_path)
    except:
        pass
    
    # 7. Pie
    c.setFont("Helvetica", 6)
    c.drawCentredString(w/2, y_tot - 4.4*cm, "Representación impresa de la")
    c.drawCentredString(w/2, y_tot - 4.7*cm, "Factura/Boleta Electrónica")
    c.drawCentredString(w/2, y_tot - 5.1*cm, "Consulte su comprobante en:")
    c.drawCentredString(w/2, y_tot - 5.4*cm, "https://play.google.com/store/apps/details?id=com.ppd.se&hl=es_US")
    c.drawCentredString(w/2, y_tot - 5.9*cm, "Anthophila: Sembrando hoy el aprendizaje de mañana.")
    c.drawCentredString(w/2, y_tot - 6.2*cm, "Gracias por su confianza.")
    
    c.save()
    return filename

def generar_mensaje_whatsapp(datos_comp):
    """Genera un mensaje de texto formateado para WhatsApp"""
    emoji_comp = "📄" if "Boleta" in datos_comp['tipo'] else "🏢"
    msg = (
        f"*{datos_comp['tipo'].upper()} ELECTRÓNICA*\n"
        f"------------------------------------------\n"
        f"✨ *ANTHOPHILA E-LEARNING EDUCATION E.I.R.L.*\n"
        f"RUC: {SUNAT_RUC}\n"
        f"Dirección: Asoc. Para Grande Mza. B Lote 15\n"
        f"Contacto: +51 906 598 622\n"
        f"------------------------------------------\n"
        f"{emoji_comp} *Serie:* {datos_comp['serie']}-{datos_comp['correlativo']}\n"
        f"📅 *Fecha:* {datos_comp['fecha']}\n"
        f"👤 *Cliente:* {datos_comp['cliente']}\n"
        f"📝 *Servicio:* {datos_comp['descripcion']}\n"
        f"💰 *Total:* S/. {datos_comp['total']:,.2f}\n"
        f"------------------------------------------\n"
        f"📎 *Adjunto envío mi comprobante en PDF.*\n"
        f"✅ *¡Gracias por confiar en nosotros!*\n"
        f"Atención especializada en neurodesarrollo."
    )
    return msg

def verificar_disponibilidad(fecha_hora_str, citas_existentes):
    nueva_cita_start = datetime.strptime(fecha_hora_str, '%Y-%m-%d %H:%M')
    nueva_cita_end = nueva_cita_start + timedelta(hours=1)
    
    # 1. Verificar Horario de Atención
    h_inicio_am = nueva_cita_start.replace(hour=9, minute=0, second=0, microsecond=0)
    h_fin_am = nueva_cita_start.replace(hour=13, minute=0, second=0, microsecond=0)
    h_inicio_pm = nueva_cita_start.replace(hour=14, minute=30, second=0, microsecond=0)
    h_fin_pm = nueva_cita_start.replace(hour=18, minute=0, second=0, microsecond=0)
    
    esta_en_horario = False
    if (nueva_cita_start >= h_inicio_am and nueva_cita_end <= h_fin_am) or \
       (nueva_cita_start >= h_inicio_pm and nueva_cita_end <= h_fin_pm):
        esta_en_horario = True
        
    if not esta_en_horario:
        return False, "Fuera del horario de atención (9am-1pm o 2:30pm-6pm)."

    # 2. Verificar Colisiones con otras citas
    for cita in citas_existentes:
        try:
            ex_start = datetime.fromisoformat(cita['start'])
            ex_end = ex_start + timedelta(hours=1)
            
            # Hay solapamiento si (S1 < E2) y (E1 > S2)
            if nueva_cita_start < ex_end and nueva_cita_end > ex_start:
                return False, f"El horario ya está ocupado por otra cita."
        except:
            continue
            
    return True, ""

# Asegurar carpetas
for folder in [BASE_DIR, IMG_FOLDER, COMPROBANTES_FOLDER]:
    if not os.path.exists(folder): os.makedirs(folder)

# --- SISTEMA DE LOGUEO ---
if 'authenticated' not in st.session_state:
    st.session_state['authenticated'] = False
    st.session_state['role'] = None

if not st.session_state['authenticated']:
    col_l, col_c, col_r = st.columns([1, 1.5, 1])
    with col_c:
        st.markdown("<br><br>", unsafe_allow_html=True)
        with st.container(border=True):
            # Cargar y mostrar el logo oficial
            logo_path = os.path.join(IMG_FOLDER, LOGO_FILE)
            if os.path.exists(logo_path):
                # Centrar la imagen usando columnas
                col_img_l, col_img_c, col_img_r = st.columns([1, 2, 1])
                with col_img_c:
                    st.image(Image.open(logo_path), width=150)
            else:
                st.markdown("<h1 style='text-align: center;'>Anthophila</h1>", unsafe_allow_html=True)
            
            st.markdown("<p style='text-align: center; color: gray;'>Sistema de Gestión de Historias Clínicas</p>", unsafe_allow_html=True)
            
            tab_admin, tab_fam = st.tabs(["🔒 Especialista", "📅 Familias"])
            with tab_admin:
                u = st.text_input("Usuario", key="user_input")
                p = st.text_input("Contraseña", type="password", key="p_admin")
                if st.button("Ingresar al Portal"):
                    if u == USUARIO_ADMIN and p == PASSWORD_ADMIN:
                        st.session_state['authenticated'] = True
                        st.session_state['role'] = 'admin'; st.rerun()
                    else:
                        st.error("Credenciales no válidas")
            with tab_fam:
                pf = st.text_input("Credencial Familiar", type="password", key="p_fam")
                if st.button("Ingresar a Citas"):
                    if pf == PASSWORD_FAMILIA:
                        st.session_state['authenticated'] = True
                        st.session_state['role'] = 'familia'; st.rerun()
                    else:
                        st.error("Credencial incorrecta")
    st.stop()

# --- INTERFAZ PRINCIPAL ---
with st.sidebar:
    logo_path = os.path.join(IMG_FOLDER, LOGO_FILE)
    if os.path.exists(logo_path):
        st.image(Image.open(logo_path), width='stretch')
    else:
        st.markdown(f"<h2 style='text-align: center;'>🐝 Anthophila</h2>", unsafe_allow_html=True)
    
    st.markdown(f"<p style='text-align: center;'><b>Rol:</b> {st.session_state['role'].capitalize()}</p>", unsafe_allow_html=True)
    st.markdown("---")
    
    if st.session_state['role'] == 'admin':
        opcion = st.radio("Menú Principal", 
                         ["🏠 Dashboard",
                          "📂 1. Apertura de Expediente", 
                          "📊 5. Agenda y Base de Datos", 
                          "📝 7. Historias Clínicas", 
                          "📋 8. Asignación de Pruebas",
                          "💰 9. Gestión de Ingresos",
                          "� 10. Facturación Electrónica",
                          "�️ 6. Mantenimiento"])
    else:
        opcion = st.radio("Menú Familiar", ["📅 Reserva de Cita", "📝 Módulo de Evaluación"])
    
    st.markdown("---")
    if st.button("🚪 Cerrar Sesión"):
        st.session_state.clear(); st.rerun()

# --- MÓDULO DASHBOARD (NUEVO) ---
if "Dashboard" in opcion:
    st.header("🏠 Dashboard de Bienvenida")
    
    # Estadísticas rápidas
    num_pacientes = len([f for f in os.listdir(BASE_DIR) if os.path.isdir(os.path.join(BASE_DIR, f))]) if os.path.exists(BASE_DIR) else 0
    citas_hoy = cargar_json(CITAS_FILE)
    num_citas = len(citas_hoy)
    
    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric("Total Pacientes", num_pacientes)
    with col2:
        st.metric("Citas Programadas", num_citas)
    with col3:
        st.metric("Informes Generados", "Pendiente") # Podría contarse luego
        
    st.markdown("---")
    st.subheader("🚀 Accesos Rápidos")
    c_r1, c_r2, c_r3 = st.columns(3)
    if c_r1.button("📂 Crear Nuevo Paciente"): 
        st.session_state['menu_opcion'] = "📂 1. Apertura de Expediente" # Nota: requiere lógica adicional para cambiar radio
    
    st.info("Bienvenido, Psic. José Miguel. Aquí puede gestionar su clínica de manera integral.")

# --- MÓDULO 1: APERTURA ---
elif "1. Apertura de Expediente" in opcion:
    st.header("📂 1. Apertura de Expediente")
    with st.container(border=True):
        st.subheader("Datos Básicos del Menor")
        with st.form("apertura_expediente"):
            c1, c2 = st.columns(2)
            nom = c1.text_input("Nombres y Apellidos del Menor", placeholder="Ej: Juan Pérez")
            dni = c2.text_input("DNI", placeholder="8 dígitos")
            edad = c1.text_input("Edad")
            padres = c2.text_input("Nombres de los Padres")
            
            st.markdown("---")
            st.subheader("Información de Contacto")
            c3, c4 = st.columns(2)
            celular = c3.text_input("Celular (WhatsApp)")
            correo = c4.text_input("Correo Electrónico")
            direccion = st.text_input("Dirección de Residencia")
            
            if st.form_submit_button("🚀 CREAR EXPEDIENTE FÍSICO"):
                if nom and dni:
                    folder_name = f"{dni}_{nom.replace(' ', '_')}"
                    path_paciente = os.path.join(BASE_DIR, folder_name)
                    os.makedirs(path_paciente, exist_ok=True)
                    
                    ficha_path = os.path.join(path_paciente, "Ficha_Personal.csv")
                    df_ficha = pd.DataFrame([{
                        "DNI": dni, "Nombre": nom, "Edad": edad, "Padres": padres,
                        "Celular": celular, "Correo": correo, "Direccion": direccion,
                        "Fecha_Apertura": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                    }])
                    df_ficha.to_csv(ficha_path, index=False, sep=';', encoding='utf-8-sig')
                    
                    st.session_state['last_dni'] = dni
                    st.session_state['last_nom'] = nom
                    st.success(f"✅ ¡Expediente de {nom} creado con éxito!")
                else:
                    st.error("⚠️ Nombre y DNI son obligatorios.")

# --- MÓDULO RESERVA: FAMILIAS ---
elif "Reserva de Cita" in opcion:
    st.header("📆 Reserva de Citas Anthophila")
    if 'temp_fechas' not in st.session_state: st.session_state.temp_fechas = []

    with st.container(border=True):
        c1, c2 = st.columns(2)
        n_nino = c1.text_input("Nombre del Paciente")
        d_nino = c2.text_input("DNI")
        n_padre = c1.text_input("Nombre del Apoderado")
        cel_padre = c2.text_input("WhatsApp (Ej: 51922...)")
        
        st.write("---")
        # Dividir en dos columnas principales para títulos y contenido
        col_main_mod, col_main_pago = st.columns(2)
        
        with col_main_mod:
            st.subheader("⚙️ Configuración de Sesiones")
            tipo_sesion = st.radio("Seleccione modalidad:", 
                                   ["Sesión Única / Evaluación", 
                                    "Pack 8 Sesiones (2 veces por semana)", 
                                    "Pack 12 Sesiones (3 veces por semana)",
                                    "Sesión Virtual"])
            
            # Mover programación de fecha y hora aquí para que esté más cerca
            st.write("---")
            st.markdown("📅 **Programar Fecha y Hora:**")
            f_c, h_c = st.columns(2)
            f_sel = f_c.date_input("Fecha de Inicio", datetime.now() + timedelta(days=1), key="f_fam")
            h_sel = h_c.time_input("Hora de la Cita", key="h_fam")
            
            if st.button("➕ AGREGAR SESIÓN A LA LISTA", width='stretch'):
                nueva_fec_hora_str = f"{f_sel} {h_sel.strftime('%H:%M')}"
                
                # Cargar citas actuales para validar
                citas_actuales = cargar_json(CITAS_FILE)
                disponible, mensaje = verificar_disponibilidad(nueva_fec_hora_str, citas_actuales + [{"start": f.replace(" ", "T")} for f in st.session_state.temp_fechas])
                
                if disponible:
                    if nueva_fec_hora_str not in st.session_state.temp_fechas:
                        st.session_state.temp_fechas.append(nueva_fec_hora_str)
                        st.rerun()
                    else:
                        st.warning("Esa fecha y hora ya está en tu lista.")
                else:
                    st.error(f"❌ {mensaje}")
        
        with col_main_pago:
            st.subheader("💳 Medios de Pago")
            
            # Sub-columnas para mostrar logos y detalles al mismo tiempo
            c_yape, c_tacna = st.columns(2)
            
            # Tamaño estandarizado para logos (Ajustado para que se vean del mismo tamaño visual)
            LOGO_WIDTH = 85
            
            with c_yape:
                # Buscar en raíz directamente (Ruta relativa para la web)
                path_y = INFO_PAGOS["Yape"]["logo"]
                if not os.path.exists(path_y):
                    path_y = os.path.join(IMG_FOLDER, INFO_PAGOS["Yape"]["logo"])
                
                if os.path.exists(path_y):
                    st.image(Image.open(path_y), width=LOGO_WIDTH)
                st.caption(f"**YAPE**\n{INFO_PAGOS['Yape']['detalle']}")
                
            with c_tacna:
                # Buscar en raíz directamente (Ruta relativa para la web)
                path_t = INFO_PAGOS["Caja Tacna"]["logo"]
                if not os.path.exists(path_t):
                    path_t = os.path.join(IMG_FOLDER, INFO_PAGOS["Caja Tacna"]["logo"])
                
                if os.path.exists(path_t):
                    # Forzamos el mismo ancho para ambos para mantener la simetría
                    st.image(Image.open(path_t), width=LOGO_WIDTH)
                st.caption(f"**CAJA TACNA**\n{INFO_PAGOS['Caja Tacna']['detalle']}")
            
            # En lugar de selectbox, asumimos que usará cualquiera de los dos
            medio_pago = "Yape / Caja Tacna"
            st.info("💡 Por favor, realice el pago por cualquiera de estos medios y adjunte el comprobante por WhatsApp.")

    if st.session_state.temp_fechas:
        st.subheader("Sesiones a registrar:")
        
        # Determinar cuántas sesiones se esperan según el tipo
        num_esperado = 1
        if "Pack 8" in tipo_sesion: num_esperado = 8
        elif "Pack 12" in tipo_sesion: num_esperado = 12
        
        num_actual = len(st.session_state.temp_fechas)
        
        if num_actual < num_esperado:
            st.warning(f"⚠️ Has seleccionado {num_actual} de {num_esperado} sesiones. Recuerda generar a la brevedad las demás para completar tu pack.")
        elif num_actual > num_esperado and num_esperado > 1:
            st.error(f"⚠️ Has seleccionado {num_actual} sesiones, pero el pack es de {num_esperado}. Por favor elimina las excedentes.")

        for i, f in enumerate(st.session_state.temp_fechas):
            col_inf, col_mod, col_del = st.columns([4, 1, 1])
            col_inf.info(f"Sesión {i+1}: {f}")
            
            # Botón Modificar
            if col_mod.button("✏️", key=f"mod_{i}"):
                st.session_state[f"editing_{i}"] = True
            
            # Botón Eliminar
            if col_del.button("🗑️", key=f"del_{i}"):
                st.session_state.temp_fechas.pop(i)
                st.rerun()
            
            # Formulario de edición si está en modo edición
            if st.session_state.get(f"editing_{i}"):
                with st.form(f"edit_form_{i}"):
                    st.write(f"Editando Sesión {i+1}")
                    # Extraer fecha y hora actual para el formulario
                    f_part, h_part = f.split(" ")
                    f_dt = datetime.strptime(f_part, "%Y-%m-%d")
                    h_dt = datetime.strptime(h_part, "%H:%M").time()
                    
                    e_fec = st.date_input("Nueva Fecha", value=f_dt)
                    e_hor = st.time_input("Nueva Hora", value=h_dt)
                    
                    if st.form_submit_button("✅ Actualizar"):
                        st.session_state.temp_fechas[i] = f"{e_fec} {e_hor.strftime('%H:%M')}"
                        del st.session_state[f"editing_{i}"]
                        st.rerun()
                    if st.form_submit_button("❌ Cancelar"):
                        del st.session_state[f"editing_{i}"]
                        st.rerun()
        
        if st.button("🚀 FINALIZAR RESERVA"):
            if not n_nino or not d_nino or not n_padre or not cel_padre:
                st.error("⚠️ Todos los datos del paciente y apoderado son obligatorios para finalizar la reserva.")
            elif num_actual == 0:
                st.error("⚠️ Debes programar al menos una sesión.")
            else:
                citas_json = cargar_json(CITAS_FILE)
                for f in st.session_state.temp_fechas:
                    citas_json.append({
                        "title": f"Cita: {n_nino} ({n_padre} - {cel_padre})", 
                        "start": f.replace(" ","T"), 
                        "color": "#27AE60"
                    })
                guardar_json(CITAS_FILE, citas_json)
                
                registrar_en_excel({"DNI": d_nino, "Paciente": n_nino, "Apoderado": n_padre, "Tipo": tipo_sesion, "Sesiones": ", ".join(st.session_state.temp_fechas)})
                
                # Generar Ticket ID
                mes_dia = datetime.now().strftime("%m%d")
                iniciales = "".join([n[0].upper() for n in n_nino.split()])[:4]
                id_reserva = f"ANT-{mes_dia}-{iniciales}"
                ahora = datetime.now().strftime("%H:%M, %d/%m/%Y")
                
                # Mensaje para el ESPECIALISTA (Lo que el padre envía)
                msg_al_especialista = (f"ID: {id_reserva}\n"
                                       f"Paciente: {n_nino}\n"
                                       f"Padre: {n_padre}\n"
                                       f"WhatsApp: {cel_padre}\n"
                                       f"Modalidad: {tipo_sesion}\n"
                                       f"Medio de Pago: {medio_pago}\n"
                                       f"Fechas: {', '.join(st.session_state.temp_fechas)}")
                
                # Mensaje de RESPUESTA (Lo que el especialista devuelve al padre)
                msg_respuesta_especialista = (f"¡Hola! Confirmo la reserva de {n_nino}. ID: {id_reserva}\n\n"
                                              f"Cita Confirmada para {n_nino}. Lo atenderá el especialista: {PSICOLOGO_FULL}. "
                                              f"Verificación C.Ps.P: https://www.cpsp.pe/busquedas/busqueda_colegiados.html#")
                
                # El link de WhatsApp enviará el mensaje AL especialista, pero estructurado para que él solo reenvíe la confirmación
                # Según tu instrucción, el padre envía la información y el especialista responde.
                # Para facilitar esto, el mensaje que se envía por WA llevará la estructura de la cita primero.
                
                url_wa = f"https://wa.me/{CELULAR_ESPECIALISTA}?text={urllib.parse.quote(msg_al_especialista + '\n\n--- RESPUESTA PARA CONFIRMAR ---\n' + msg_respuesta_especialista)}"
                
                st.success(f"✅ ¡Reserva guardada con éxito!")
                
                # --- TICKET VISUAL FÍSICO ---
                st.markdown(f"""
                <div style="border: 2px dashed #27AE60; padding: 20px; border-radius: 15px; background-color: #f9f9f9; color: black; font-family: 'Courier New', Courier, monospace;">
                    <h2 style="text-align: center; color: #27AE60; margin-top: 0;">🎟️ TICKET DE RESERVA</h2>
                    <p style="text-align: center; font-weight: bold; font-size: 1.2rem;">ID: {id_reserva}</p>
                    <hr style="border-top: 1px solid #ccc;">
                    <p><b>PACIENTE:</b> {n_nino}</p>
                    <p><b>APODERADO:</b> {n_padre}</p>
                    <p><b>WHATSAPP:</b> {cel_padre}</p>
                    <p><b>MODALIDAD:</b> {tipo_sesion}</p>
                    <p><b>MEDIO DE PAGO:</b> {medio_pago}</p>
                    <p><b>FECHAS:</b></p>
                    <p style="font-size: 0.9rem;">{', '.join(st.session_state.temp_fechas)}</p>
                    <hr style="border-top: 1px dashed #ccc;">
                    <p style="text-align: center; font-size: 0.8rem; color: gray;">Generado el {ahora}</p>
                </div>
                """, unsafe_allow_html=True)

                with st.container(border=True):
                    st.subheader("📢 Pasos finales para la familia")
                    st.write("1. **Confirmar con el Especialista:** Haz clic abajo para enviar el mensaje de confirmación listo para el Psicólogo.")
                    st.markdown(f'''<a href="{url_wa}" target="_blank"><button style="background-color:#25D366;color:white;width:100%;padding:15px;border:none;border-radius:10px;font-weight:bold;cursor:pointer;font-size:1rem;">✉️ ENVIAR CONFIRMACIÓN POR WHATSAPP</button></a>''', unsafe_allow_html=True)
                    
                    st.write("---")
                    st.write("2. **Recordatorio Personal:** Descarga el calendario para activar alertas automáticas en tu celular y correo.")
                    
                    # Generar archivo ICS (iCalendar) compatible con Google, Outlook, Apple
                    ics_content = "BEGIN:VCALENDAR\nVERSION:2.0\nPRODID:-//Anthophila//ES\n"
                    for f_str in st.session_state.temp_fechas:
                        dt = datetime.strptime(f_str, '%Y-%m-%d %H:%M')
                        f_start = dt.strftime('%Y%m%dT%H%M%S')
                        f_end = (dt + timedelta(hours=1)).strftime('%Y%m%dT%H%M%S')
                        ics_content += "BEGIN:VEVENT\n"
                        ics_content += f"SUMMARY:Cita Anthophila: {n_nino}\n"
                        ics_content += f"DTSTART:{f_start}\n"
                        ics_content += f"DTEND:{f_end}\n"
                        ics_content += f"DESCRIPTION:Cita confirmada con el Especialista José Miguel Granda. ID: {id_reserva}\n"
                        ics_content += "END:VEVENT\n"
                    ics_content += "END:VCALENDAR"
                    
                    st.download_button(
                        label="📅 DESCARGAR CALENDARIO (Para Google/Outlook)",
                        data=ics_content,
                        file_name=f"citas_anthophila_{id_reserva}.ics",
                        mime="text/calendar",
                        width='stretch'
                    )
                    st.caption("ℹ️ Abre este archivo para agregar las citas automáticamente a tu calendario personal.")

                st.session_state.temp_fechas = []

# --- MÓDULO DE EVALUACIÓN: FAMILIAS ---
elif "Módulo de Evaluación" in opcion:
    st.header("📝 Módulo de Evaluación Familiar")
    
    if 'dni_hijo_auth' not in st.session_state:
        st.session_state['dni_hijo_auth'] = None
    
    if not st.session_state['dni_hijo_auth']:
        with st.container(border=True):
            st.subheader("Acceso Protegido")
            st.write("Por favor, ingrese el DNI de su hijo para acceder a las pruebas asignadas.")
            dni_hijo = st.text_input("DNI del Niño/a (Actúa como contraseña)", type="password")
            if st.button("Ingresar al Módulo"):
                if dni_hijo:
                    # Verificar si existe el expediente
                    folders = [f for f in os.listdir(BASE_DIR) if f.startswith(dni_hijo)]
                    if folders:
                        st.session_state['dni_hijo_auth'] = dni_hijo
                        st.session_state['folder_hijo'] = folders[0]
                        st.success("Acceso concedido.")
                        st.rerun()
                    else:
                        st.error("No se encontró un expediente con ese DNI. Por favor, verifique o contacte al especialista.")
                else:
                    st.warning("Ingrese el DNI.")
    else:
        # Ya autenticado con el DNI del hijo
        folder_paciente = st.session_state['folder_hijo']
        path_paciente = os.path.join(BASE_DIR, folder_paciente)
        nom_paciente = folder_paciente.split("_", 1)[1].replace("_", " ")
        
        st.success(f"🔓 Sesión iniciada para: **{nom_paciente}**")
        if st.button("Cerrar Módulo de Evaluación"):
            st.session_state['dni_hijo_auth'] = None
            st.rerun()
            
        st.markdown("---")
        st.subheader("📋 Pruebas Psicológicas Disponibles")
        
        # Lista de pruebas (esto se puede expandir)
        pruebas_disponibles = ["Cuestionario de Screening Sensorial", "Inventario de Conducta (Brief)"]
        
        # Filtrar por asignación del especialista
        asignaciones = cargar_json(ASIGNACIONES_FILE)
        if st.session_state['dni_hijo_auth'] in asignaciones:
            pruebas_permitidas = asignaciones[st.session_state['dni_hijo_auth']]
            if not isinstance(pruebas_permitidas, list): pruebas_permitidas = [pruebas_permitidas]
            
            # Solo mostrar las que el especialista asignó
            lista_opciones = ["-- Seleccione --"] + [p for p in pruebas_disponibles if p in pruebas_permitidas]
            if len(lista_opciones) == 1:
                st.info("ℹ️ El especialista aún no ha activado pruebas específicas para hoy.")
        else:
            # Si no hay asignación, por defecto no mostramos ninguna para obligar a que el admin asigne
            # o mostramos todas si prefieres, pero el usuario pidió "designar que prueba debe contestar"
            st.warning("⚠️ No tiene pruebas asignadas actualmente. Por favor, consulte con el especialista.")
            lista_opciones = ["-- Seleccione --"]

        prueba_sel = st.selectbox("Seleccione la prueba a completar:", lista_opciones)
        
        if prueba_sel == "Cuestionario de Screening Sensorial":
            with st.form("test_sensorial"):
                st.write("### Cuestionario de Screening Sensorial")
                st.write("Responda con qué frecuencia observa estas conductas en su hijo/a.")
                
                q1 = st.select_slider("1. ¿Parece excesivamente sensible a ruidos fuertes?", options=["Nunca", "Rara vez", "A veces", "Frecuentemente", "Siempre"])
                q2 = st.select_slider("2. ¿Evita texturas de ropa o alimentos específicos?", options=["Nunca", "Rara vez", "A veces", "Frecuentemente", "Siempre"])
                q3 = st.select_slider("3. ¿Busca constantemente movimiento o dar vueltas?", options=["Nunca", "Rara vez", "A veces", "Frecuentemente", "Siempre"])
                q4 = st.select_slider("4. ¿Se distrae fácilmente ante estímulos visuales?", options=["Nunca", "Rara vez", "A veces", "Frecuentemente", "Siempre"])
                
                observaciones = st.text_area("Observaciones adicionales del padre/madre")
                
                if st.form_submit_button("Enviar Respuestas"):
                    respuestas = {
                        "Prueba": "Screening Sensorial",
                        "Fecha": datetime.now().strftime("%Y-%m-%d %H:%M"),
                        "R1_Ruido": q1,
                        "R2_Texturas": q2,
                        "R3_Movimiento": q3,
                        "R4_Visual": q4,
                        "Observaciones": observaciones
                    }
                    df_resp = pd.DataFrame([respuestas])
                    filename = f"Prueba_Sensorial_{datetime.now().strftime('%Y%m%d_%H%M')}.csv"
                    df_resp.to_csv(os.path.join(path_paciente, filename), index=False, sep=';', encoding='utf-8-sig')
                    st.success("✅ ¡Muchas gracias! Sus respuestas han sido enviadas directamente a la historia clínica de su hijo/a.")

        elif prueba_sel == "Inventario de Conducta (Brief)":
            with st.form("test_conducta"):
                st.write("### Inventario de Conducta")
                q1 = st.radio("¿Tiene dificultades para seguir instrucciones?", ["Sí", "No", "A veces"])
                q2 = st.radio("¿Muestra cambios bruscos de humor?", ["Sí", "No", "A veces"])
                q3 = st.radio("¿Le cuesta esperar su turno?", ["Sí", "No", "A veces"])
                
                if st.form_submit_button("Enviar Respuestas"):
                    respuestas = {
                        "Prueba": "Inventario Conducta",
                        "Fecha": datetime.now().strftime("%Y-%m-%d %H:%M"),
                        "R1_Instrucciones": q1,
                        "R2_Humor": q2,
                        "R3_Turno": q3
                    }
                    df_resp = pd.DataFrame([respuestas])
                    filename = f"Prueba_Conducta_{datetime.now().strftime('%Y%m%d_%H%M')}.csv"
                    df_resp.to_csv(os.path.join(path_paciente, filename), index=False, sep=';', encoding='utf-8-sig')
                    st.success("✅ ¡Respuestas guardadas correctamente!")


# --- MÓDULO 5: GESTIÓN (ADMIN) ---
elif "5. Agenda y Base de Datos" in opcion:
    st.header("📊 Gestión de Agenda")
    citas = cargar_json(CITAS_FILE)
    
    t_cal, t_db, t_add, t_edit = st.tabs(["📅 Vista Calendario", "📂 Base de Datos Excel", "➕ Nueva Cita", "✏️ Gestionar"])
    
    with t_cal: 
        with st.container(border=True):
            st.subheader("Calendario de Sesiones")
            calendar(events=citas, options={"initialView": "dayGridMonth", "headerToolbar": {"left": "prev,next today", "center": "title", "right": "dayGridMonth,timeGridWeek,timeGridDay"}}, key="cal_admin_full")
        
    with t_db:
        if os.path.exists(EXCEL_CITAS):
            df = pd.read_csv(EXCEL_CITAS, sep=';', encoding='utf-8-sig', on_bad_lines='skip')
            st.dataframe(df, width='stretch')
            st.download_button("📥 Descargar Excel Completo", data=df.to_csv(index=False, sep=';', encoding='utf-8-sig'), file_name="db_reservas_anthophila.csv", mime="text/csv")
        else:
            st.info("No hay datos registrados aún.")
            
    with t_add:
        with st.container(border=True):
            st.subheader("Agregar Cita Manual")
            with st.form("form_add_cita"):
                c1, c2 = st.columns(2)
                new_pac = c1.text_input("Paciente")
                new_dni = c2.text_input("DNI")
                new_fec = c1.date_input("Fecha")
                new_hor = c2.time_input("Hora")
                new_msg = st.text_input("Nota (Celular/Motivo)")
                
                if st.form_submit_button("📅 REGISTRAR CITA"):
                    if new_pac:
                        fec_hora_str = f"{new_fec} {new_hor.strftime('%H:%M')}"
                        disponible, mensaje = verificar_disponibilidad(fec_hora_str, citas)
                        
                        if disponible:
                            citas.append({
                                "title": f"Cita: {new_pac} ({new_msg})", 
                                "start": f"{new_fec}T{new_hor.strftime('%H:%M')}", 
                                "color": "#3498DB",
                                "extendedProps": {"dni": new_dni}
                            })
                            guardar_json(CITAS_FILE, citas)
                            registrar_en_excel({"DNI": new_dni, "Paciente": new_pac, "Tipo": "Manual Admin", "Sesiones": fec_hora_str})
                            st.success(f"✅ Cita registrada para {new_pac}.")
                            st.rerun()
                        else:
                            st.error(f"❌ {mensaje}")

    with t_edit:
        if citas:
            lista_citas = [f"{i} | {c['title']} | {c['start']}" for i, c in enumerate(citas)]
            cita_sel = st.selectbox("Seleccione cita para gestionar:", lista_citas)
            idx = int(cita_sel.split(" | ")[0])
            cita_actual = citas[idx]
            
            st.write("---")
            st.subheader("✏️ Modificar Cita")
            with st.form(f"form_mod_cita_{idx}"):
                c1, c2 = st.columns(2)
                nuevo_titulo = c1.text_input("Título de la Cita", value=cita_actual['title'])
                nueva_fec_hora = c2.text_input("Fecha y Hora (ISO)", value=cita_actual['start'], help="Formato: YYYY-MM-DDTHH:MM")
                
                if st.form_submit_button("💾 GUARDAR CAMBIOS"):
                    old_start = cita_actual['start']
                    citas[idx]['title'] = nuevo_titulo
                    citas[idx]['start'] = nueva_fec_hora
                    guardar_json(CITAS_FILE, citas)
                    
                    # Sincronizar con Excel si existe
                    if os.path.exists(EXCEL_CITAS):
                        df = pd.read_csv(EXCEL_CITAS, sep=';', encoding='utf-8-sig')
                        mask = df['Sesiones'].astype(str).str.contains(old_start.replace("T", " "))
                        if mask.any():
                            df.loc[mask, 'Sesiones'] = nueva_fec_hora.replace("T", " ")
                            df.loc[mask, 'Paciente'] = nuevo_titulo.replace("Cita: ", "").split(" (")[0]
                            df.to_csv(EXCEL_CITAS, index=False, sep=';', encoding='utf-8-sig')
                    
                    st.success("✅ Cita actualizada correctamente en Calendario y Excel.")
                    st.rerun()
            
            st.write("---")
            if st.button("🗑️ ELIMINAR SELECCIÓN DEFINITIVAMENTE", type="primary"):
                cita_a_borrar = citas.pop(idx)
                guardar_json(CITAS_FILE, citas)
                
                # Sincronizar con Excel si existe
                if os.path.exists(EXCEL_CITAS):
                    df = pd.read_csv(EXCEL_CITAS, sep=';', encoding='utf-8-sig')
                    fec_hora_busqueda = cita_a_borrar['start'].replace("T", " ")
                    df = df[~df['Sesiones'].astype(str).str.contains(fec_hora_busqueda)]
                    df.to_csv(EXCEL_CITAS, index=False, sep=';', encoding='utf-8-sig')
                
                st.success("✅ Cita eliminada de Calendario y Excel.")
                st.rerun()

# --- MÓDULO 7: HISTORIAS CLÍNICAS ---
elif "7. Historias Clínicas" in opcion:
    st.header("📝 Gestión de Historias Clínicas")
    
    with st.container(border=True):
        st.subheader("🔍 Localización de Expediente")
        dni_input = st.text_input("Ingrese DNI del Paciente", value=st.session_state.get('last_dni', ''), placeholder="Ej: 12345678")
    
    if dni_input:
        folders = [f for f in os.listdir(BASE_DIR) if f.startswith(dni_input)]
        if folders:
            folder_paciente = folders[0]
            nom_paciente = folder_paciente.replace(f"{dni_input}_", "").replace("_", " ")
            st.success(f"✅ **Paciente:** {nom_paciente}")
            
            # --- ALERTA DE DEUDA (NUEVO) ---
            if os.path.exists(INGRESOS_FILE):
                df_aux = pd.read_csv(INGRESOS_FILE, sep=';', encoding='utf-8-sig')
                if 'DNI' in df_aux.columns:
                    df_pac_deuda = df_aux[(df_aux['DNI'].astype(str) == str(dni_input)) & (df_aux['Estado'].isin(['PENDIENTE', 'DEBE']))]
                    if not df_pac_deuda.empty:
                        total_deuda = pd.to_numeric(df_pac_deuda['Monto'], errors='coerce').sum()
                        st.error(f"⚠️ **ATENCIÓN:** El paciente tiene una deuda pendiente de **S/. {total_deuda:,.2f}**")
                        with st.expander("Ver detalle de deuda"):
                            st.table(df_pac_deuda[["Fecha", "Monto", "Observaciones"]])

            path_paciente = os.path.join(BASE_DIR, folder_paciente)
            
            # --- SELECCIÓN DE FORMATO ---
            st.markdown("---")
            st.subheader("📂 Formatos y Herramientas")
            
            c_sel1, c_sel2 = st.columns([1, 2])
            with c_sel1:
                opc_formato = st.radio("Elija una acción:", 
                                      ["📄 Ver Archivos / Subir", 
                                       "👶 Informe Infantil (Auto)", 
                                       "🧠 Historia Breve", 
                                       "🏃 Sesión de Trabajo", 
                                       "📓 Bitácora Clínica"])
            
            with c_sel2:
                # --- LÓGICA DE CADA OPCIÓN ---
                if opc_formato == "📄 Ver Archivos / Subir":
                    st.info("Gestión de archivos del paciente")
                    uploaded_file = st.file_uploader("Subir PDF, Imagen o Word", type=['pdf', 'jpg', 'png', 'docx', 'doc'])
                    if uploaded_file:
                        with open(os.path.join(path_paciente, uploaded_file.name), "wb") as f:
                            f.write(uploaded_file.getbuffer())
                        st.success("Archivo subido.")
                    
                    st.write("### Archivos en Carpeta:")
                    archivos = os.listdir(path_paciente)
                    for a in archivos:
                        st.text(f"• {a}")

                elif opc_formato == "👶 Informe Infantil (Auto)":
                    with st.form("form_infantil_v2"):
                        st.write("### 📋 Formato Digital para Generar Informes Automáticos")
                        st.write("#### 1. Datos del Niño")
                        c1, c2, c3 = st.columns(3)
                        i_nom = c1.text_input("Nombre", value=nom_paciente)
                        i_edad = c2.text_input("Edad")
                        i_fec_nac = c3.date_input("Fecha de nacimiento")
                        i_col = c1.text_input("Colegio")
                        i_grado = c2.text_input("Grado")
                        
                        st.write("#### 2. Motivo de Consulta")
                        i_motivo = st.text_area("Describa las principales preocupaciones...")
                        
                        st.write("#### 3. Desarrollo y Lenguaje")
                        c_d1, c_d2 = st.columns(2)
                        i_camino = c_d1.radio("¿A qué edad caminó?", ["antes de 12 meses", "12-18 meses", "después de 18 meses"], horizontal=True)
                        i_habla = c_d2.radio("¿El niño?", ["habla frases completas", "usa palabras sueltas", "no habla"], horizontal=True)
                        
                        st.write("#### 4. Interacción y Conducta")
                        i_social = st.radio("Interacción Social", ["juega con otros niños", "prefiere jugar solo", "evita interacción"], horizontal=True)
                        
                        st.write("#### 5. Aprendizaje Escolar")
                        c_ap1, c_ap2, c_ap3, c_ap4 = st.columns(4)
                        d_lec = c_ap1.checkbox("Lectura")
                        d_esc = c_ap2.checkbox("Escritura")
                        d_mat = c_ap3.checkbox("Matemática")
                        d_ate = c_ap4.checkbox("Atención")
                        
                        st.write("#### 6. Conducta")
                        c_co1, c_co2, c_co3, c_co4 = st.columns(4)
                        d_hip = c_co1.checkbox("Hiperactividad")
                        d_rab = c_co2.checkbox("Rabietas")
                        d_rep = c_co3.checkbox("Conductas repetitivas")
                        d_sen = c_co4.checkbox("Sensibilidad sensorial")
                        
                        i_obs = st.text_area("Observación del Psicólogo")
                        
                        if st.form_submit_button("🚀 GENERAR INFORME WORD"):
                            inf_data = {
                                "Paciente": i_nom, "DNI": dni_input, "Edad": i_edad, "Colegio": i_col, "Grado": i_grado,
                                "Motivo": i_motivo, "Camino": i_camino, "Lenguaje": i_habla, "Social": i_social,
                                "Dificultades": f"{'Lectura ' if d_lec else ''}{'Escritura ' if d_esc else ''}{'Matemática ' if d_mat else ''}{'Atención' if d_ate else ''}",
                                "Conducta": f"{'Hiperactividad ' if d_hip else ''}{'Rabietas ' if d_rab else ''}{'Repetitivas ' if d_rep else ''}{'Sensorial' if d_sen else ''}",
                                "Obs": i_obs
                            }
                            docx_name = os.path.join(path_paciente, f"Informe_{datetime.now().strftime('%Y%m%d')}.docx")
                            exportar_a_docx(inf_data, docx_name)
                            st.success("Informe generado en la carpeta.")

                elif opc_formato == "🧠 Historia Breve":
                    with st.form("form_breve_v2"):
                        st.write("### 🧠 Historia Clínica Psicológica Breve (Consultorio)")
                        st.write("#### I. Datos del Paciente")
                        c1, c2 = st.columns(2)
                        b_nom = c1.text_input("Nombre", value=nom_paciente)
                        b_edad = c2.text_input("Edad")
                        b_tel = c1.text_input("Teléfono")
                        b_fec = c2.date_input("Fecha")
                        
                        b_motivo = st.text_area("Motivo de Consulta")
                        
                        st.write("#### II. Historia del Problema")
                        col1, col2, col3 = st.columns(3)
                        b_inicio = col1.text_input("Inicio")
                        b_desencadenantes = col2.text_input("Situaciones desencadenantes")
                        b_impacto = col3.text_input("Impacto en la vida diaria")
                        
                        st.write("#### III. Antecedentes")
                        a1, a2, a3 = st.columns(3)
                        b_ant_med = a1.text_area("Médicos")
                        b_ant_psic = a2.text_area("Psicológicos")
                        b_ant_fam = a3.text_area("Familiares")
                        
                        st.write("#### IV. Evaluación Inicial")
                        e1, e2 = st.columns(2)
                        b_eval_emo = e1.text_input("Estado emocional")
                        b_eval_con = e2.text_input("Conducta")
                        b_eval_cog = e1.text_input("Cognición")
                        b_eval_soc = e2.text_input("Relaciones sociales")
                        
                        b_diag = st.text_input("Diagnóstico Presuntivo")
                        
                        st.write("#### V. Objetivos Terapéuticos")
                        obj_col1, obj_col2 = st.columns([1, 2])
                        b_obj = obj_col1.text_input("Objetivo")
                        b_est = obj_col2.text_input("Estrategia")
                        
                        st.write("#### VI. Seguimiento de Sesiones")
                        seg_col1, seg_col2 = st.columns([1, 3])
                        b_seg_fec = seg_col1.date_input("Fecha Sesión")
                        b_seg_obs = seg_col2.text_area("Observación")
                        
                        if st.form_submit_button("💾 GUARDAR HISTORIA BREVE"):
                            breve_data = {
                                "Nombre": b_nom, "DNI": dni_input, "Edad": b_edad, "Telefono": b_tel, "Fecha": str(b_fec),
                                "Motivo": b_motivo, "Inicio": b_inicio, "Desencadenantes": b_desencadenantes, "Impacto": b_impacto,
                                "Ant_Med": b_ant_med, "Ant_Psic": b_ant_psic, "Ant_Fam": b_ant_fam,
                                "Eval_Emo": b_eval_emo, "Eval_Con": b_eval_con, "Eval_Cog": b_eval_cog, "Eval_Soc": b_eval_soc,
                                "Diagnostico": b_diag, "Objetivo": b_obj, "Estrategia": b_est, "Seg_Fecha": str(b_seg_fec), "Seg_Obs": b_seg_obs
                            }
                            pd.DataFrame([breve_data]).to_csv(os.path.join(path_paciente, f"HCL_Breve_{datetime.now().strftime('%Y%m%d')}.csv"), index=False, sep=';', encoding='utf-8-sig')
                            st.success("Guardado.")

                elif opc_formato == "🏃 Sesión de Trabajo":
                    with st.form("form_sesion_v2"):
                        st.write("### 🏃 Sesión de Trabajo Terapéutico")
                        st.write("#### 1. Datos de Control")
                        c1, c2 = st.columns(2)
                        s_num = c1.text_input("N° de Sesión (Ej: 4/12 del Pack)")
                        s_estado = c2.selectbox("Estado Inicial (Soporte y Calma)", ["Regulado", "Ansioso", "Apático", "Irritable"])
                        
                        st.write("#### 2. Ejecución del Plan (DUA)")
                        s_obj = st.text_input("Objetivo Terapéutico")
                        c_act1, c_act2 = st.columns(2)
                        s_prep = c_act1.text_area("Preparación Motriz")
                        s_cent = c_act2.text_area("Actividad Central")
                        s_apoyo = st.radio("Nivel de Apoyo", ["Autónomo", "Apoyo Verbal", "Apoyo Físico", "Modelamiento"], horizontal=True)
                        
                        st.write("#### 3. Perfil de Desempeño (Barreras y Logros)")
                        c_per1, c_per2 = st.columns(2)
                        s_aten = c_per1.text_area("Atención y Enfoque (B4?)")
                        s_cond = c_per2.text_area("Regulación Conductual (Frustración)")
                        s_logro = st.text_input("Logro Significativo (Hoy logró...)")
                        
                        st.write("#### 4. Pendientes y Tarea (Anthophila en Casa)")
                        s_padre = st.text_area("Sugerencia al padre")
                        s_prox = st.text_area("Próximo paso")
                        
                        if st.form_submit_button("💾 REGISTRAR PROGRESO"):
                            ses_data = {
                                "Paciente": nom_paciente, "DNI": dni_input, "Num_Sesion": s_num, "Estado": s_estado,
                                "Objetivo": s_obj, "Prep_Motriz": s_prep, "Act_Central": s_cent, "Apoyo": s_apoyo,
                                "Atencion": s_aten, "Conducta": s_cond, "Logro": s_logro, "Sugerencia": s_padre, "Proximo": s_prox
                            }
                            docx_name = os.path.join(path_paciente, f"Sesion_{s_num.replace('/','_')}.docx")
                            exportar_a_docx(ses_data, docx_name)
                            st.success("Sesión guardada.")

                elif opc_formato == "📓 Bitácora Clínica":
                    with st.form("form_bitacora_v2"):
                        st.write("### 📓 Bitácora de Sesión Clínica (Uso Interno)")
                        st.write("#### I. Identificación")
                        c1, c2 = st.columns(2)
                        b_ses = c1.text_input("Sesión N° (Ej: ____ / 12)")
                        b_hor = c2.time_input("Hora", value=datetime.now().time())
                        b_fec = st.date_input("Fecha", value=datetime.now())
                        
                        st.write("#### II. Registro de Proceso")
                        b_cont = st.text_area("Describa el proceso...", height=250)
                        
                        if st.form_submit_button("💾 GUARDAR BITÁCORA"):
                            f_bit = os.path.join(path_paciente, f"Bitacora_{datetime.now().strftime('%Y%m%d_%H%M')}.txt")
                            with open(f_bit, "w", encoding='utf-8') as f:
                                f.write(f"Bitácora de {nom_paciente} - Sesión {b_ses}\n")
                                f.write(f"Fecha: {b_fec} | Hora: {b_hor}\n\n{b_cont}")
                            st.success("Bitácora actualizada.")
        else:
            st.error("❌ No existe expediente con ese DNI.")
    else:
        st.warning("Ingrese un DNI para comenzar.")

# --- MÓDULO 8: ASIGNACIÓN DE PRUEBAS (ADMIN) ---
elif "8. Asignación de Pruebas" in opcion:
    st.header("📋 Asignación de Pruebas Psicológicas")
    
    with st.container(border=True):
        st.subheader("Selección de Paciente y Prueba")
        
        # Obtener lista de pacientes
        expedientes = [f for f in os.listdir(BASE_DIR) if os.path.isdir(os.path.join(BASE_DIR, f))]
        
        if expedientes:
            c1, c2 = st.columns(2)
            exp_sel = c1.selectbox("Seleccione el Paciente:", ["-- Seleccione --"] + expedientes)
            
            pruebas_list = ["Cuestionario de Screening Sensorial", "Inventario de Conducta (Brief)"]
            pruebas_sel = c2.multiselect("Pruebas a habilitar para la familia:", pruebas_list)
            
            if st.button("💾 GUARDAR ASIGNACIÓN"):
                if exp_sel != "-- Seleccione --" and pruebas_sel:
                    dni_pac = exp_sel.split("_")[0]
                    asignaciones = cargar_json(ASIGNACIONES_FILE)
                    asignaciones[dni_pac] = pruebas_sel
                    guardar_json(ASIGNACIONES_FILE, asignaciones)
                    st.success(f"✅ Se han habilitado {len(pruebas_sel)} pruebas para el DNI: {dni_pac}")
                else:
                    st.error("⚠️ Debe seleccionar un paciente y al menos una prueba.")
        else:
            st.info("No hay pacientes registrados aún.")
            
    st.markdown("---")
    st.subheader("📋 Asignaciones Actuales")
    asignaciones_actuales = cargar_json(ASIGNACIONES_FILE)
    if asignaciones_actuales:
        for d, p in asignaciones_actuales.items():
            col_d, col_p, col_b = st.columns([1, 2, 1])
            col_d.write(f"**DNI:** {d}")
            col_p.write(f"**Pruebas:** {', '.join(p)}")
            if col_b.button("🗑️ Quitar", key=f"del_asig_{d}"):
                del asignaciones_actuales[d]
                guardar_json(ASIGNACIONES_FILE, asignaciones_actuales)
                st.rerun()
    else:
        st.write("No hay asignaciones activas.")

# --- MÓDULO 9: GESTIÓN DE INGRESOS (ADMIN) ---
elif "9. Gestión de Ingresos" in opcion:
    st.header("💰 Gestión de Ingresos y Contabilidad")
    
    # --- ALERTA SUNAT (RUC termina en 7) ---
    cronograma_ruc7 = {
        1: "2026-02-20", 2: "2026-03-20", 3: "2026-04-23", 4: "2026-05-22",
        5: "2026-06-19", 6: "2026-07-21", 7: "2026-08-24", 8: "2026-09-21",
        9: "2026-10-22", 10: "2026-11-20", 11: "2026-12-23", 12: "2027-01-22"
    }
    
    hoy = datetime.now()
    mes_actual = hoy.month
    vencimiento_str = cronograma_ruc7.get(mes_actual)
    if vencimiento_str:
        vencimiento_dt = datetime.strptime(vencimiento_str, "%Y-%m-%d")
        dias_faltantes = (vencimiento_dt - hoy).days
        
        if 0 <= dias_faltantes <= 7:
            st.warning(f"⚠️ **ALERTA SUNAT (RUC 7):** El vencimiento del periodo actual es el **{vencimiento_dt.strftime('%d/%m/%Y')}**. Faltan {dias_faltantes} días.")
        elif dias_faltantes < 0:
            # Buscar el del próximo mes si el de este mes ya pasó
            proximo_mes = mes_actual + 1 if mes_actual < 12 else 1
            vencimiento_str = cronograma_ruc7.get(proximo_mes)
            vencimiento_dt = datetime.strptime(vencimiento_str, "%Y-%m-%d")
            st.info(f"📅 **PRÓXIMO VENCIMIENTO SUNAT (RUC 7):** {vencimiento_dt.strftime('%d/%m/%Y')}")
        else:
            st.info(f"📅 **VENCIMIENTO SUNAT (RUC 7):** {vencimiento_dt.strftime('%d/%m/%Y')} (Faltan {dias_faltantes} días)")

    # Cargar datos
    if os.path.exists(INGRESOS_FILE):
        df_ing = pd.read_csv(INGRESOS_FILE, sep=';', encoding='utf-8-sig', dtype={'DNI': str, 'RUC': str})
    else:
        df_ing = pd.DataFrame(columns=["Fecha", "Estado", "Monto", "Cliente", "DNI", "RUC", "Medio", "Destino", "Observaciones", "Salidas"])
    
    # Asegurar tipos numéricos para cálculos
    df_ing['Monto'] = pd.to_numeric(df_ing['Monto'], errors='coerce').fillna(0.0).astype(float)
    df_ing['Salidas'] = pd.to_numeric(df_ing['Salidas'], errors='coerce').fillna(0.0).astype(float)
    df_ing['DNI'] = df_ing['DNI'].astype(str).replace('nan', '')
    df_ing['RUC'] = df_ing['RUC'].astype(str).replace('nan', '')

    tab1, tab2, tab3, tab4, tab5 = st.tabs(["📊 Registro", "📉 Contabilidad", "👥 Por Paciente", "📋 Métricas", "📖 Libros SUNAT"])
    
    with tab1:
        with st.container(border=True):
            st.subheader("➕ Registrar Nuevo Movimiento")
            with st.form("form_ingresos"):
                c1, c2, c3 = st.columns(3)
                fec = c1.date_input("Fecha", datetime.now())
                est = c2.selectbox("Estado", ["PAGADO", "PENDIENTE", "SALIDA", "DEBE"])
                mon = c3.number_input("Monto (S/.)", min_value=0.0, step=0.1)
                
                c4, c5, c6 = st.columns(3)
                cli = c4.text_input("Cliente / Concepto")
                dni_mov = c5.text_input("DNI (Opcional)", placeholder="8 dígitos")
                ruc_prov = c6.text_input("RUC Proveedor (Solo para Salidas)", placeholder="11 dígitos")
                
                c7, c8, c9 = st.columns(3)
                med = c7.selectbox("Medio de Pago", ["", "EFECTIVO", "YAPE", "TRANSFERENCIA", "TARJETA"])
                des = c8.selectbox("Destino", ["", "CAJA", "BANCO"])
                sal = c9.number_input("Salidas / Gastos (S/.)", min_value=0.0, step=0.1)
                
                obs = st.text_area("Observaciones")
                
                if st.form_submit_button("💾 GUARDAR MOVIMIENTO"):
                    nuevo_mov = {
                        "Fecha": fec.strftime("%Y-%m-%d"),
                        "Estado": est,
                        "Monto": mon,
                        "Cliente": cli,
                        "DNI": dni_mov,
                        "RUC": ruc_prov,
                        "Medio": med,
                        "Destino": des,
                        "Observaciones": obs,
                        "Salidas": sal
                    }
                    df_ing = pd.concat([df_ing, pd.DataFrame([nuevo_mov])], ignore_index=True)
                    df_ing.to_csv(INGRESOS_FILE, index=False, sep=';', encoding='utf-8-sig')
                    
                    # --- AUTO GENERAR ASIENTO CONTABLE ---
                    monto_contable = mon if est != "SALIDA" else sal
                    es_gasto = (est == "SALIDA")
                    generar_asiento_contable(fec.strftime("%Y-%m-%d"), monto_contable, cli, obs, es_salida=es_gasto)
                    
                    # --- AUTO GENERAR REGISTRO DE COMPRA (Si es salida) ---
                    if es_gasto:
                        generar_registro_compra(fec.strftime("%Y-%m-%d"), sal, cli, ruc_prov, obs)
                    else:
                        # --- AUTO GENERAR REGISTRO DE VENTA (Si es ingreso pagado) ---
                        if est == "PAGADO":
                            generar_registro_venta(fec.strftime("%Y-%m-%d"), mon, cli, dni_mov, obs)
                    
                    st.success("✅ Movimiento, Asiento y Registro Contable guardados.")
                    st.rerun()

        st.markdown("---")
        st.subheader("📋 Historial de Movimientos")
        st.dataframe(df_ing, width='stretch')
        
        if not df_ing.empty:
            if st.button("🗑️ Eliminar último registro"):
                df_ing = df_ing[:-1]
                df_ing.to_csv(INGRESOS_FILE, index=False, sep=';', encoding='utf-8-sig')
                st.rerun()

    with tab2:
        st.subheader("📉 Desglose Contable (Impuestos, Ahorro, Sueldo)")
        if not df_ing.empty:
            # Filtrar solo los pagados para el análisis contable
            df_pagados = df_ing[df_ing['Estado'] == 'PAGADO'].copy()
            
            if not df_pagados.empty:
                # Conversión explícita a float64 para evitar FutureWarning
                df_pagados['Monto'] = pd.to_numeric(df_pagados['Monto'], errors='coerce').astype('float64').fillna(0.0)
                
                # Cálculos según la lógica proporcionada
                df_pagados['Impuestos (20%)'] = (df_pagados['Monto'] * 0.20).astype('float64')
                df_pagados['Ahorros (10%)'] = (df_pagados['Monto'] * 0.10).astype('float64')
                df_pagados['Contingencia (10%)'] = (df_pagados['Monto'] * 0.10).astype('float64')
                df_pagados['Sueldo (60%)'] = (df_pagados['Monto'] * 0.60).astype('float64')
                
                # Acumulados
                df_pagados['Acum. Impuestos'] = df_pagados['Impuestos (20%)'].cumsum().astype('float64')
                df_pagados['Acum. Ahorro'] = df_pagados['Ahorros (10%)'].cumsum().astype('float64')
                df_pagados['Acum. Contingencia'] = df_pagados['Contingencia (10%)'].cumsum().astype('float64')
                df_pagados['Acum. Sueldo'] = df_pagados['Sueldo (60%)'].cumsum().astype('float64')
                
                # Mostrar tabla con columnas relevantes
                cols_view = ["Fecha", "Cliente", "Monto", "Impuestos (20%)", "Ahorros (10%)", "Contingencia (10%)", "Sueldo (60%)", 
                             "Acum. Impuestos", "Acum. Ahorro", "Acum. Contingencia", "Acum. Sueldo"]
                st.dataframe(df_pagados[cols_view], width='stretch')
            else:
                st.info("No hay ingresos 'PAGADOS' para mostrar en el desglose.")
        else:
            st.info("No hay datos registrados.")

    with tab3:
        st.subheader("👥 Estado de Cuentas por Paciente")
        if not df_ing.empty:
            df_ing['Monto'] = pd.to_numeric(df_ing['Monto'], errors='coerce').fillna(0)
            
            # Obtener lista de pacientes con DNI en ingresos
            pacientes_con_dni = df_ing[df_ing['DNI'].notna() & (df_ing['DNI'] != "")]['DNI'].unique()
            
            if len(pacientes_con_dni) > 0:
                dni_sel = st.selectbox("Seleccione DNI del Paciente:", pacientes_con_dni)
                
                if dni_sel:
                    df_pac = df_ing[df_ing['DNI'] == dni_sel].copy()
                    
                    # Totales
                    total_pagado = df_pac[df_pac['Estado'] == 'PAGADO']['Monto'].sum()
                    total_debe = df_pac[df_pac['Estado'].isin(['PENDIENTE', 'DEBE'])]['Monto'].sum()
                    
                    c1, c2 = st.columns(2)
                    c1.metric("✅ Total Pagado", f"S/. {total_pagado:,.2f}")
                    c2.metric("🔴 Total Deuda Actual", f"S/. {total_debe:,.2f}")
                    
                    st.write("#### Detalle de Movimientos del Paciente")
                    st.dataframe(df_pac[["Fecha", "Estado", "Monto", "Cliente", "Medio", "Observaciones"]], width='stretch')
            else:
                st.info("No hay movimientos vinculados a un DNI aún.")
        else:
            st.info("No hay datos registrados.")

    with tab4:
        st.subheader("📊 Resumen de Métricas")
        if not df_ing.empty:
            df_ing['Monto'] = pd.to_numeric(df_ing['Monto'], errors='coerce').fillna(0)
            df_ing['Salidas'] = pd.to_numeric(df_ing['Salidas'], errors='coerce').fillna(0)
            
            total_cobrado = df_ing[df_ing['Estado'] == 'PAGADO']['Monto'].sum()
            total_pendiente = df_ing[df_ing['Estado'].isin(['PENDIENTE', 'DEBE'])]['Monto'].sum()
            total_salidas = df_ing['Salidas'].sum()
            
            # Resumen por medio de pago
            yape_total = df_ing[(df_ing['Medio'] == 'YAPE') & (df_ing['Estado'] == 'PAGADO')]['Monto'].sum()
            efectivo_total = df_ing[(df_ing['Medio'] == 'EFECTIVO') & (df_ing['Estado'] == 'PAGADO')]['Monto'].sum()
            
            c1, c2, c3 = st.columns(3)
            c1.metric("💰 Total Cobrado", f"S/. {total_cobrado:,.2f}")
            c2.metric("⏳ Total Pendiente", f"S/. {total_pendiente:,.2f}")
            c3.metric("💸 Total Salidas", f"S/. {total_salidas:,.2f}")
            
            st.markdown("---")
            c4, c5 = st.columns(2)
            c4.metric("📱 Pagos por Yape", f"S/. {yape_total:,.2f}")
            c5.metric("💵 Pagos en Efectivo", f"S/. {efectivo_total:,.2f}")
            
            # Gráfico simple (opcional con streamlit)
            st.write("#### Balance Final (Cobrado - Salidas)")
            balance = total_cobrado - total_salidas
            st.title(f"S/. {balance:,.2f}")
        else:
            st.info("No hay datos para generar métricas.")

    with tab5:
        st.subheader("📖 Libros Electrónicos SUNAT (PLE)")
        st.write("Generación de formatos 5.2, 5.4, 8.1, 8.2 y 14.1.")
        
        # --- BOTÓN DE SINCRONIZACIÓN (NUEVO) ---
        if st.button("🔄 Sincronizar/Generar Libros desde Movimientos", width='stretch'):
            if not df_ing.empty:
                # Limpiar archivos actuales para regenerar
                for f in [LIBRO_DIARIO_FILE, REGISTRO_COMPRAS_8_1_FILE, REGISTRO_VENTAS_14_1_FILE]:
                    if os.path.exists(f): os.remove(f)
                
                # Procesar cada fila de ingresos
                for _, row in df_ing.iterrows():
                    m_fec = row['Fecha']
                    m_mon = pd.to_numeric(row['Monto'], errors='coerce') or 0.0
                    m_sal = pd.to_numeric(row['Salidas'], errors='coerce') or 0.0
                    m_cli = row['Cliente']
                    m_est = row['Estado']
                    m_dni = row.get('DNI', "")
                    m_ruc = row.get('RUC', "")
                    m_obs = row['Observaciones']
                    
                    monto_contable = m_mon if m_est != "SALIDA" else m_sal
                    es_gasto = (m_est == "SALIDA")
                    
                    # Generar asiento
                    generar_asiento_contable(m_fec, monto_contable, m_cli, m_obs, es_salida=es_gasto)
                    
                    # Generar registros auxiliares
                    if es_gasto:
                        generar_registro_compra(m_fec, m_sal, m_cli, m_ruc, m_obs)
                    elif m_est == "PAGADO":
                        generar_registro_venta(m_fec, m_mon, m_cli, m_dni, m_obs)
                
                st.success("✅ ¡Libros contables sincronizados con éxito!")
                st.rerun()
            else:
                st.warning("No hay movimientos registrados para sincronizar.")

        st.markdown("---")
        c_p1, c_p2, c_p3 = st.columns(3)
        
        with c_p1:
            st.write("### 📗 Libro Diario Simplificado 5.2")
            if os.path.exists(LIBRO_DIARIO_FILE):
                df_diario = pd.read_csv(LIBRO_DIARIO_FILE, sep=';', encoding='utf-8-sig')
                st.dataframe(df_diario.head(5), width='stretch')
                st.download_button(
                    label="📥 Descargar Libro Diario Simplificado 5.2", 
                    data=df_diario.to_csv(index=False, sep='|', encoding='utf-8'), 
                    file_name=f"LE{datetime.now().strftime('%Y%m')}00050200001111.txt"
                )
            else:
                st.info("ℹ️ Registre movimientos de ingresos o salidas para generar el Libro Diario.")
            
            st.write("### 📘 Plan Contable 5.4")
            plan_data = [
                {"Periodo": datetime.now().strftime("%Y0101"), "Cuenta": "1212", "Desc": "Cuentas por cobrar", "Plan": "01", "Estado": "1"},
                {"Periodo": datetime.now().strftime("%Y0101"), "Cuenta": "7011", "Desc": "Ventas de servicios", "Plan": "01", "Estado": "1"},
                {"Periodo": datetime.now().strftime("%Y0101"), "Cuenta": "6391", "Desc": "Gastos servicios", "Plan": "01", "Estado": "1"},
                {"Periodo": datetime.now().strftime("%Y0101"), "Cuenta": "4212", "Desc": "Emitidas", "Plan": "01", "Estado": "1"}
            ]
            df_plan = pd.DataFrame(plan_data)
            st.dataframe(df_plan, width='stretch')
            st.download_button(label="📥 Descargar Plan 5.4", data=df_plan.to_csv(index=False, sep='|', encoding='utf-8'), file_name=f"LE{datetime.now().strftime('%Y%m')}00050400001111.txt")

        with c_p2:
            st.write("### 📙 Registro de Compras 8.1")
            if os.path.exists(REGISTRO_COMPRAS_8_1_FILE):
                df_c81 = pd.read_csv(REGISTRO_COMPRAS_8_1_FILE, sep=';', encoding='utf-8-sig')
                st.dataframe(df_c81.head(5), width='stretch')
                st.download_button(label="📥 Descargar Compras 8.1", data=df_c81.to_csv(index=False, sep='|', encoding='utf-8'), file_name=f"LE{datetime.now().strftime('%Y%m')}00080100001111.txt")
            else:
                st.info("No hay compras registradas.")

            st.write("### 📕 Registro de Compras 8.2")
            st.info("Formato para no domiciliados.")
            st.download_button(label="📥 Descargar Compras 8.2 (Vacío)", data="", file_name=f"LE{datetime.now().strftime('%Y%m')}00080200001111.txt")

        with c_p3:
            st.write("### 📘 Registro de Ventas 14.1")
            if os.path.exists(REGISTRO_VENTAS_14_1_FILE):
                df_v141 = pd.read_csv(REGISTRO_VENTAS_14_1_FILE, sep=';', encoding='utf-8-sig')
                st.dataframe(df_v141.head(5), width='stretch')
                st.download_button(label="📥 Descargar Ventas 14.1", data=df_v141.to_csv(index=False, sep='|', encoding='utf-8'), file_name=f"LE{datetime.now().strftime('%Y%m')}00140100001111.txt")
            else:
                st.info("No hay ventas registradas.")
            
            st.write("### 📓 Ventas Simplificado 14.2")
            st.info("Este libro se deriva del 14.1 para regímenes simplificados.")
            st.download_button(label="📥 Descargar Ventas 14.2 (Vacío)", data="", file_name=f"LE{datetime.now().strftime('%Y%m')}00140200001111.txt")

# --- MÓDULO 10: FACTURACIÓN ELECTRÓNICA ---
elif "10. Facturación Electrónica" in opcion:
    st.header("📄 Facturación Electrónica SUNAT")
    
    # Cargar base de datos de clientes
    if os.path.exists(DB_CLIENTES_FILE):
        df_clientes = pd.read_csv(DB_CLIENTES_FILE, sep=';', encoding='utf-8-sig', dtype={'DNI_RUC': str, 'Nombre': str, 'Direccion': str})
    else:
        df_clientes = pd.DataFrame(columns=["DNI_RUC", "Nombre", "Direccion"])
    
    # Asegurar tipos explícitos para evitar FutureWarnings
    df_clientes['DNI_RUC'] = df_clientes['DNI_RUC'].astype(str)
    df_clientes['Nombre'] = df_clientes['Nombre'].astype(str)
    df_clientes['Direccion'] = df_clientes['Direccion'].astype(str)

    with st.expander("⚙️ Configuración SUNAT (RUC / Clave SOL / Certificado)"):
        c_ruc, c_user, c_pass = st.columns(3)
        ruc_emp = c_ruc.text_input("Número de RUC", value=SUNAT_RUC)
        user_sol = c_user.text_input("Usuario SOL", value=SUNAT_USER)
        pass_sol = c_pass.text_input("Clave SOL", type="password", value=SUNAT_PASS)
        cert_file = st.file_uploader("Cargar Certificado Digital (.pfx)", type=["pfx"])
        if st.button("Guardar Configuración"):
            st.success("Configuración guardada localmente.")

    tab_emitir, tab_historial = st.tabs(["🆕 Emitir Comprobante", "📜 Historial de Comprobantes"])
    
    with tab_emitir:
        with st.container(border=True):
            st.subheader("Datos del Comprobante")
            col1, col2, col3 = st.columns(3)
            tipo_comp = col1.selectbox("Tipo de Comprobante", ["Boleta de Venta", "Factura", "Nota de Crédito"])
            
            # Lógica para Nota de Crédito
            es_nota_credito = tipo_comp == "Nota de Crédito"
            
            if es_nota_credito:
                serie_def = "FC01" # Serie estándar para notas de crédito vinculadas a Factura
                correlativo_def = "000001"
            else:
                serie_def = "B001" if tipo_comp == "Boleta de Venta" else "F001"
                correlativo_def = "000001"

            serie = col2.text_input("Serie", value=serie_def)
            correlativo = col3.text_input("Correlativo", value=correlativo_def)
            
            if es_nota_credito:
                st.info("⚠️ Las Notas de Crédito deben estar vinculadas a un comprobante emitido previamente.")
                c_nc1, c_nc2 = st.columns(2)
                doc_ref = c_nc1.text_input("Documento de Referencia (Ej: F001-000123)")
                motivo_nc = c_nc2.selectbox("Motivo (Catálogo SUNAT No. 09)", MOTIVOS_NOTA_CREDITO)
            else:
                doc_ref = ""
                motivo_nc = ""

            st.markdown("---")
            st.subheader("Datos del Cliente")
            col4, col5 = st.columns(2)
            doc_cliente = col4.text_input("DNI / RUC del Cliente", key="doc_cliente_input")
            
            # Lógica de autocompletado de cliente
            nom_cliente_sugerido = ""
            dir_cliente_sugerido = ""
            if doc_cliente:
                cliente_match = df_clientes[df_clientes['DNI_RUC'].astype(str) == str(doc_cliente)]
                if not cliente_match.empty:
                    nom_cliente_sugerido = cliente_match.iloc[0]['Nombre']
                    dir_cliente_sugerido = cliente_match.iloc[0]['Direccion']
                    st.info(f"✨ Cliente encontrado: {nom_cliente_sugerido}")

            nom_cliente = col5.text_input("Nombre / Razón Social", value=nom_cliente_sugerido)
            dir_cliente = st.text_input("Dirección (Opcional)", value=dir_cliente_sugerido)
            grabar_cliente = st.checkbox("💾 Grabar/Actualizar datos de este cliente", value=True)
            
            st.markdown("---")
            st.subheader("Detalle del Servicio")
            col6, col7, col8 = st.columns([2, 1, 1])
            
            # Selector de servicios
            servicio_pre = col6.selectbox("Seleccione Servicio:", SERVICIOS_DISPONIBLES)
            if servicio_pre == "Otro (Especificar)":
                desc_serv = st.text_input("Especifique el servicio:", placeholder="Ej: Terapia Ocupacional")
            else:
                desc_serv = servicio_pre

            cant = col7.number_input("Cantidad", min_value=1, value=1)
            precio_unit = col8.number_input("Precio Unitario (S/.)", min_value=0.0, value=0.0)
            
            total = cant * precio_unit
            st.write(f"### Total a Cobrar: S/. {total:,.2f}")
            
            # --- DATOS DEL COMPROBANTE ---
            datos_comp = {
                "tipo": tipo_comp,
                "serie": serie,
                "correlativo": correlativo,
                "cliente": nom_cliente,
                "doc_cliente": doc_cliente,
                "fecha": datetime.now().strftime("%Y-%m-%d"),
                "direccion": dir_cliente,
                "descripcion": desc_serv,
                "cantidad": cant,
                "total": total,
                "doc_referencia": doc_ref,
                "motivo_nc": motivo_nc
            }

            col_p1, col_p2 = st.columns(2)
            
            with col_p1:
                if st.button("👁️ PREVISUALIZAR", width='stretch'):
                    if not doc_cliente or not nom_cliente or total <= 0:
                        st.error("Complete los datos para previsualizar.")
                    else:
                        st.markdown("### 📋 Vista Previa del Comprobante")
                        with st.container(border=True):
                            c1_pre, c2_pre = st.columns([1, 2])
                            logo_path = os.path.join(IMG_FOLDER, LOGO_FILE)
                            if os.path.exists(logo_path):
                                c1_pre.image(Image.open(logo_path), width='stretch')
                            
                            c2_pre.markdown(f"""
                                **ANTHOPHILA E-LEARNING EDUCATION E.I.R.L.**  
                                *Atención especializada en desarrollo cognitivo y aprendizaje*  
                                RUC: {SUNAT_RUC}  
                                Dirección: Asoc. Para Grande Mza. B Lote 15  
                                Contacto: +51 906 598 622
                            """)
                            
                            st.markdown("---")
                            col_a, col_b = st.columns(2)
                            with col_a:
                                st.write(f"**Documento:** {tipo_comp}")
                                st.write(f"**Serie-Correlativo:** {serie}-{correlativo}")
                                st.write(f"**Fecha:** {datos_comp['fecha']}")
                            with col_b:
                                st.write(f"**Señor(es):** {nom_cliente}")
                                st.write(f"**DNI/RUC:** {doc_cliente}")
                                st.write(f"**Moneda:** Soles (PEN)")
                            
                            st.markdown("---")
                            st.write(f"**Detalle del Servicio:** {desc_serv}")
                            st.write(f"**Cantidad:** {cant}")
                            st.write(f"**Total a Pagar:** S/. {total:,.2f}")
                            st.info("💡 Esto es solo una vista previa. Use el botón 'EMITIR' para generar el PDF legal.")

            with col_p2:
                if st.button("🚀 EMITIR Y ENVIAR A SUNAT", width='stretch'):
                    if not doc_cliente or not nom_cliente or total <= 0:
                        st.error("Por favor complete los datos obligatorios.")
                    else:
                        with st.spinner("Generando Comprobante..."):
                            # 1. Guardar/Actualizar Cliente si se solicitó
                            if grabar_cliente:
                                # Convertimos a string explícitamente antes de comparar/asignar
                                doc_str = str(doc_cliente)
                                nuevo_cliente = pd.DataFrame([{"DNI_RUC": doc_str, "Nombre": str(nom_cliente), "Direccion": str(dir_cliente)}])
                                
                                if not df_clientes[df_clientes['DNI_RUC'] == doc_str].empty:
                                    df_clientes.loc[df_clientes['DNI_RUC'] == doc_str, ['Nombre', 'Direccion']] = [str(nom_cliente), str(dir_cliente)]
                                else:
                                    df_clientes = pd.concat([df_clientes, nuevo_cliente], ignore_index=True)
                                
                                df_clientes.to_csv(DB_CLIENTES_FILE, index=False, sep=';', encoding='utf-8-sig')

                            # 2. Generar PDFs
                            pdf_filename = f"{tipo_comp.replace(' ', '_')}_{serie}_{correlativo}.pdf"
                            pdf_path = os.path.join(COMPROBANTES_FOLDER, pdf_filename)
                            result_pdf = generar_pdf_comprobante(datos_comp, pdf_path)
                            
                            ticket_filename = f"Ticket_{serie}_{correlativo}.pdf"
                            ticket_path = os.path.join(COMPROBANTES_FOLDER, ticket_filename)
                            generar_pdf_ticket_termico(datos_comp, ticket_path)
                            
                            msg_wa = generar_mensaje_whatsapp(datos_comp)
                            
                            # 3. Registrar en Ingresos y Contabilidad
                            # Si es Nota de Crédito, el monto es negativo en ingresos (devolución)
                            monto_registro = -total if es_nota_credito else total
                            
                            nuevo_mov = {
                                "Fecha": datos_comp['fecha'], "Estado": "PAGADO", "Monto": monto_registro,
                                "Cliente": nom_cliente, "DNI": doc_cliente if len(doc_cliente) == 8 else "",
                                "RUC": doc_cliente if len(doc_cliente) == 11 else "", "Medio": "TRANSFERENCIA",
                                "Destino": "BANCO", "Observaciones": f"{tipo_comp} {serie}-{correlativo} {doc_ref}", "Salidas": 0
                            }
                            if os.path.exists(INGRESOS_FILE):
                                df_ing_tmp = pd.read_csv(INGRESOS_FILE, sep=';', encoding='utf-8-sig')
                            else:
                                df_ing_tmp = pd.DataFrame(columns=["Fecha", "Estado", "Monto", "Cliente", "DNI", "RUC", "Medio", "Destino", "Observaciones", "Salidas"])
                            
                            df_ing_tmp = pd.concat([df_ing_tmp, pd.DataFrame([nuevo_mov])], ignore_index=True)
                            df_ing_tmp.to_csv(INGRESOS_FILE, index=False, sep=';', encoding='utf-8-sig')
                            
                            generar_asiento_contable(nuevo_mov['Fecha'], total, nom_cliente, nuevo_mov['Observaciones'])
                            generar_registro_venta(nuevo_mov['Fecha'], total, nom_cliente, doc_cliente, nuevo_mov['Observaciones'])
                            
                            st.success(f"✅ {tipo_comp} emitida correctamente.")
                            st.balloons()
                            
                            # Botones de acción
                            c_d1, c_d2, c_d3 = st.columns(3)
                            if result_pdf and os.path.exists(pdf_path):
                                with open(pdf_path, "rb") as f:
                                    c_d1.download_button(label="📥 FORMATO A5 (PDF)", data=f, file_name=pdf_filename, mime="application/pdf", width='stretch')
                            
                            if os.path.exists(ticket_path):
                                with open(ticket_path, "rb") as f:
                                    c_d2.download_button(label="📠 TICKET 80mm (PDF)", data=f, file_name=ticket_filename, mime="application/pdf", width='stretch')
                            
                            wa_url = f"https://wa.me/?text={urllib.parse.quote(msg_wa)}"
                            c_d3.markdown(f'''<a href="{wa_url}" target="_blank"><button style="background-color:#25D366;color:white;width:100%;padding:10px;border:none;border-radius:8px;font-weight:bold;cursor:pointer;">📱 WHATSAPP</button></a>''', unsafe_allow_html=True)
                            
                            st.info("✅ **¡Listo!** Puedes descargar los archivos arriba o enviarlos por WhatsApp.")
                            # st.rerun()  <-- Eliminado para que los botones de descarga no desaparezcan al instante

    with tab_historial:
        st.subheader("📜 Historial de Comprobantes Electrónicos")
        if os.path.exists(REGISTRO_VENTAS_14_1_FILE):
            df_ventas = pd.read_csv(REGISTRO_VENTAS_14_1_FILE, sep=';', encoding='utf-8-sig')
            
            # Mostrar tabla
            st.dataframe(df_ventas, width='stretch')
            
            st.markdown("---")
            st.subheader("📥 Descargar Comprobantes Guardados")
            archivos_pdf = [f for f in os.listdir(COMPROBANTES_FOLDER) if f.endswith(".pdf")]
            
            if archivos_pdf:
                col_arch, col_btn = st.columns([3, 1])
                archivo_sel = col_arch.selectbox("Seleccione un comprobante para descargar:", archivos_pdf)
                
                if archivo_sel:
                    with open(os.path.join(COMPROBANTES_FOLDER, archivo_sel), "rb") as f:
                        col_btn.download_button(
                            label="📥 DESCARGAR SELECCIONADO",
                            data=f,
                            file_name=archivo_sel,
                            mime="application/pdf",
                            width='stretch'
                        )
            else:
                st.info("No se encontraron archivos PDF en la carpeta de comprobantes.")
        else:
            st.info("No hay comprobantes emitidos aún.")

# --- MÓDULO 6: MANTENIMIENTO ---
elif "6. Mantenimiento" in opcion:
    st.header("🛠️ Herramientas de Limpieza y Pruebas")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.subheader("Bases de Datos")
        if st.button("🧹 ELIMINAR TODAS LAS CITAS (Excel y Calendario)", width='stretch'):
            if os.path.exists(EXCEL_CITAS): os.remove(EXCEL_CITAS)
            if os.path.exists(CITAS_FILE): os.remove(CITAS_FILE)
            st.success("Toda la agenda (Excel y Calendario) ha sido eliminada.")
            st.rerun()
    
    with col2:
        st.subheader("Gestión de Expedientes")
        # Listar carpetas de pacientes existentes
        if os.path.exists(BASE_DIR):
            expedientes = [f for f in os.listdir(BASE_DIR) if os.path.isdir(os.path.join(BASE_DIR, f))]
            if expedientes:
                exp_borrar = st.selectbox("Seleccione expediente a eliminar:", ["-- Seleccione --"] + expedientes)
                
                if exp_borrar != "-- Seleccione --":
                    st.warning(f"⚠️ Esta acción eliminará permanentemente la carpeta: **{exp_borrar}** y todos sus archivos.")
                    confirmar = st.checkbox(f"Confirmo que deseo borrar el expediente de {exp_borrar}")
                    
                    if st.button("🔥 BORRAR EXPEDIENTE", type="primary", disabled=not confirmar, width='stretch'):
                        import shutil
                        shutil.rmtree(os.path.join(BASE_DIR, exp_borrar))
                        st.success(f"Expediente {exp_borrar} eliminado correctamente.")
                        st.rerun()
            else:
                st.info("No hay expedientes creados.")
        else:
            st.error("No se encontró la carpeta base de pacientes.")